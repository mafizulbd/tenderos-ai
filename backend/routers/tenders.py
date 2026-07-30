"""Tender routes: list/detail/update, analyze (sync + streaming), reanalyze,
delete, DOCX/PDF export, AI proposal wizard, AI bid strategy."""

import asyncio
import threading
from datetime import datetime
from io import BytesIO

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import StreamingResponse
from sqlalchemy import or_
from sqlalchemy.orm import Session

from database import SessionLocal
from deps import (
    MAX_UPLOAD_BYTES, _can_modify_tender, _company_profile, _extract_bid_score,
    _get_knowledge_base, _get_org, _parse_deadline, _sse, check_usage_limit,
    extract_text_from_file, get_current_membership, get_current_user, get_db,
    get_tender_response, increment_usage, limiter,
)
from hermes_client import (
    analyze_with_gemini, parse_gemini_response, stream_bid_strategy,
    stream_personalized_proposal, stream_with_gemini,
)
from models import OrgMembership, Tender, User
from schemas import ProposalWizardRequest, ReanalyzeRequest, TenderUpdate

router = APIRouter()

# ---------------------------------------------------------------------------
# Tender routes
# ---------------------------------------------------------------------------


@router.get("/tenders")
def list_tenders(
    search: str = "",
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    query = db.query(Tender).filter(Tender.organization_id == membership.organization_id)

    if search.strip():
        term = f"%{search.strip()}%"
        query = query.filter(
            or_(Tender.title.ilike(term), Tender.file_name.ilike(term))
        )

    tenders = query.order_by(Tender.id.desc()).all()
    return [
        {
            "id": t.id,
            "user_id": t.user_id,
            "title": t.title,
            "language": t.language,
            "status": t.status,
            "file_name": t.file_name or "",
            "file_size": t.file_size or 0,
            "deadline": t.deadline,
            "bid_status": t.bid_status or "reviewing",
            "bid_score": t.bid_score,
            "approval_status": t.approval_status or "none",
            "created_at": t.created_at,
            "summary": (t.summary or "")[:180],
        }
        for t in tenders
    ]


@router.get("/tenders/{tender_id}")
def get_tender(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    return get_tender_response(tender)


@router.patch("/tenders/{tender_id}")
def update_tender(
    tender_id: int,
    payload: TenderUpdate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only edit tenders you created.")

    valid_bid_statuses = {"reviewing", "submitted", "won", "lost", "no-bid"}
    if payload.bid_status is not None:
        if payload.bid_status not in valid_bid_statuses:
            raise HTTPException(status_code=400, detail=f"Invalid bid status. Must be one of: {', '.join(valid_bid_statuses)}")
        tender.bid_status = payload.bid_status

    if payload.notes is not None:
        tender.notes = payload.notes

    if payload.deadline is not None:
        tender.deadline = _parse_deadline(payload.deadline)

    db.commit()
    db.refresh(tender)
    return get_tender_response(tender)


@router.post("/tenders/analyze")
@limiter.limit("10/minute")
async def analyze_tender(
    request: Request,
    title: str = Form(...),
    language: str = Form("english"),
    deadline: str = Form(""),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    clean_title = title.strip()
    if not clean_title:
        raise HTTPException(status_code=400, detail="Tender title is required.")
    if language not in {"english", "bangla"}:
        raise HTTPException(status_code=400, detail="Unsupported language.")

    org = _get_org(membership.organization_id, db)
    check_usage_limit(org, db)

    content = await file.read()
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File is too large. Maximum upload is 15 MB.")

    tender_text = extract_text_from_file(file.filename, content)
    if len(tender_text.strip()) < 20:
        raise HTTPException(status_code=400, detail="Unable to extract readable text from this file.")

    result = analyze_with_gemini(tender_text, language, _company_profile(current_user, org))
    bid_score = _extract_bid_score(result.get("bid_recommendation"))

    tender = Tender(
        user_id=current_user.id,
        organization_id=membership.organization_id,
        title=clean_title,
        language=language,
        status="failed" if (result.get("summary") or "").startswith("Gemini API Error:") else "completed",
        file_name=file.filename,
        file_type=file.content_type or "",
        file_size=len(content),
        deadline=_parse_deadline(deadline),
        bid_score=bid_score,
        original_text=tender_text,
        summary=result.get("summary"),
        eligibility=result.get("eligibility"),
        financial_requirements=result.get("financial_requirements"),
        required_documents=result.get("required_documents"),
        compliance_matrix=result.get("compliance_matrix"),
        risk_analysis=result.get("risk_analysis"),
        bid_recommendation=result.get("bid_recommendation"),
        proposal_draft=result.get("proposal_draft"),
        final_checklist=result.get("final_checklist"),
    )
    db.add(tender)
    db.commit()
    db.refresh(tender)
    increment_usage(org, db)
    return get_tender_response(tender)


@router.post("/tenders/analyze-stream")
@limiter.limit("10/minute")
async def analyze_tender_stream(
    request: Request,
    title: str = Form(...),
    language: str = Form("english"),
    deadline: str = Form(""),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    clean_title = title.strip()
    if not clean_title:
        raise HTTPException(status_code=400, detail="Tender title is required.")
    if language not in {"english", "bangla"}:
        raise HTTPException(status_code=400, detail="Unsupported language.")

    org = _get_org(membership.organization_id, db)
    check_usage_limit(org, db)

    content = await file.read()
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File is too large. Maximum upload is 15 MB.")

    tender_text = extract_text_from_file(file.filename, content)
    if len(tender_text.strip()) < 20:
        raise HTTPException(status_code=400, detail="Unable to extract readable text from this file.")

    file_name = file.filename
    file_type = file.content_type or ""
    file_size = len(content)
    user_id = current_user.id
    organization_id = membership.organization_id
    profile = _company_profile(current_user, org)
    parsed_deadline = _parse_deadline(deadline)

    async def generate():
        yield _sse({"type": "progress", "stage": "analyzing", "message": "AI is analyzing your tender..."})

        loop = asyncio.get_running_loop()
        queue: asyncio.Queue = asyncio.Queue()

        def produce():
            try:
                for chunk in stream_with_gemini(tender_text, language, profile):
                    asyncio.run_coroutine_threadsafe(queue.put(("chunk", chunk)), loop).result()
            except Exception as exc:
                asyncio.run_coroutine_threadsafe(queue.put(("error", str(exc))), loop).result()
            finally:
                asyncio.run_coroutine_threadsafe(queue.put(("done", None)), loop).result()

        thread = threading.Thread(target=produce, daemon=True)
        thread.start()

        accumulated: list[str] = []
        while True:
            kind, data = await queue.get()
            if kind == "chunk":
                accumulated.append(data)
                yield _sse({"type": "chunk", "text": data})
            elif kind == "error":
                yield _sse({"type": "error", "detail": data})
                thread.join(timeout=5)
                return
            elif kind == "done":
                break

        thread.join(timeout=5)
        yield _sse({"type": "progress", "stage": "saving", "message": "Saving results..."})

        full_text = "".join(accumulated)
        result = parse_gemini_response(full_text)
        bid_score = _extract_bid_score(result.get("bid_recommendation"))
        status = "failed" if (result.get("summary") or "").startswith("Gemini API Error:") else "completed"

        tender = Tender(
            user_id=user_id,
            organization_id=organization_id,
            title=clean_title,
            language=language,
            status=status,
            file_name=file_name,
            file_type=file_type,
            file_size=file_size,
            deadline=parsed_deadline,
            bid_score=bid_score,
            original_text=tender_text,
            summary=result.get("summary"),
            eligibility=result.get("eligibility"),
            financial_requirements=result.get("financial_requirements"),
            required_documents=result.get("required_documents"),
            compliance_matrix=result.get("compliance_matrix"),
            risk_analysis=result.get("risk_analysis"),
            bid_recommendation=result.get("bid_recommendation"),
            proposal_draft=result.get("proposal_draft"),
            final_checklist=result.get("final_checklist"),
        )
        db.add(tender)
        db.commit()
        db.refresh(tender)
        increment_usage(org, db)

        yield _sse({"type": "done", "tender": get_tender_response(tender)})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.post("/tenders/{tender_id}/reanalyze")
@limiter.limit("10/minute")
def reanalyze_tender(
    request: Request,
    tender_id: int,
    payload: ReanalyzeRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only re-analyze tenders you created.")
    if payload.language not in {"english", "bangla"}:
        raise HTTPException(status_code=400, detail="Unsupported language.")
    if not tender.original_text:
        raise HTTPException(status_code=400, detail="Original document text not available for re-analysis.")

    org = _get_org(membership.organization_id, db)
    check_usage_limit(org, db)

    result = analyze_with_gemini(tender.original_text, payload.language, _company_profile(current_user, org))
    bid_score = _extract_bid_score(result.get("bid_recommendation"))

    tender.language = payload.language
    tender.status = "failed" if (result.get("summary") or "").startswith("Gemini API Error:") else "completed"
    tender.summary = result.get("summary")
    tender.eligibility = result.get("eligibility")
    tender.financial_requirements = result.get("financial_requirements")
    tender.required_documents = result.get("required_documents")
    tender.compliance_matrix = result.get("compliance_matrix")
    tender.risk_analysis = result.get("risk_analysis")
    tender.bid_recommendation = result.get("bid_recommendation")
    tender.proposal_draft = result.get("proposal_draft")
    tender.final_checklist = result.get("final_checklist")
    tender.bid_score = bid_score

    db.commit()
    db.refresh(tender)
    increment_usage(org, db)
    return get_tender_response(tender)


@router.delete("/tenders/{tender_id}")
def delete_tender(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only delete tenders you created.")
    db.delete(tender)
    db.commit()
    return {"detail": "Tender deleted."}


@router.get("/tenders/{tender_id}/export-docx")
def export_docx(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")

    from docx import Document

    doc = Document()
    doc.add_heading("TenderOS AI — Tender Analysis Report", 0)
    doc.add_heading(tender.title, level=1)
    doc.add_paragraph(f"Output Language: {tender.language.title()}")
    if tender.deadline:
        doc.add_paragraph(f"Submission Deadline: {tender.deadline.strftime('%d %B %Y')}")
    if tender.bid_score is not None:
        doc.add_paragraph(f"AI Bid Score: {tender.bid_score}/100")
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%d %B %Y')}")

    sections = [
        ("Executive Summary",           tender.summary),
        ("Eligibility Criteria",        tender.eligibility),
        ("Financial Requirements",      tender.financial_requirements),
        ("Required Documents",          tender.required_documents),
        ("Compliance Matrix",           tender.compliance_matrix),
        ("Risk Analysis",               tender.risk_analysis),
        ("Bid Recommendation",          tender.bid_recommendation),
        ("Tender Submission Draft",     tender.proposal_draft),
        ("Final Submission Checklist",  tender.final_checklist),
    ]
    for heading, content in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(content or "Not available")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=tender_{tender.id}_report.docx"},
    )


@router.get("/tenders/{tender_id}/export-pdf")
def export_pdf(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")

    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(120, 120, 120)
            self.cell(0, 6, "TenderOS AI — Confidential Analysis Report", align="R")
            self.ln(8)

        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", size=8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 6, f"Page {self.page_no()}", align="C")

    pdf = PDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # Title block
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(30, 60, 120)
    pdf.multi_cell(0, 10, "TenderOS AI", align="C")
    pdf.set_font("Helvetica", size=11)
    pdf.set_text_color(60, 60, 60)
    pdf.multi_cell(0, 7, "Bangladesh Procurement Analysis Report", align="C")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(20, 20, 20)

    def safe(txt: str) -> str:
        return (txt or "").encode("latin-1", errors="replace").decode("latin-1")

    pdf.multi_cell(0, 8, safe(tender.title), align="C")
    pdf.ln(2)

    pdf.set_font("Helvetica", size=9)
    pdf.set_text_color(100, 100, 100)
    meta_parts = [f"Language: {tender.language.title()}"]
    if tender.deadline:
        meta_parts.append(f"Deadline: {tender.deadline.strftime('%d %b %Y')}")
    if tender.bid_score is not None:
        meta_parts.append(f"Bid Score: {tender.bid_score}/100")
    pdf.cell(0, 6, "  |  ".join(meta_parts), align="C")
    pdf.ln(10)

    pdf.set_draw_color(200, 210, 230)
    pdf.set_line_width(0.4)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(8)
    pdf.set_text_color(0, 0, 0)

    report_sections = [
        ("Executive Summary",           tender.summary),
        ("Eligibility Criteria",        tender.eligibility),
        ("Financial Requirements",      tender.financial_requirements),
        ("Required Documents",          tender.required_documents),
        ("Compliance Matrix",           tender.compliance_matrix),
        ("Risk Analysis",               tender.risk_analysis),
        ("Bid Recommendation",          tender.bid_recommendation),
        ("Tender Submission Draft",     tender.proposal_draft),
        ("Final Submission Checklist",  tender.final_checklist),
    ]

    for heading, content in report_sections:
        if not content:
            continue
        # Section heading
        pdf.set_fill_color(235, 241, 255)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(30, 60, 120)
        pdf.multi_cell(0, 8, heading, fill=True)
        pdf.ln(2)
        # Section body
        pdf.set_font("Helvetica", size=9)
        pdf.set_text_color(30, 30, 30)
        try:
            pdf.multi_cell(0, 5.5, safe(content))
        except Exception:
            pdf.multi_cell(0, 5.5, "[Content not renderable in PDF. Use DOCX export for Bengali text.]")
        pdf.ln(6)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=tender_{tender.id}_report.pdf"},
    )


# ---------------------------------------------------------------------------
# AI Proposal Wizard — generate personalized proposal with company KB
# ---------------------------------------------------------------------------

@router.post("/tenders/{tender_id}/generate-proposal")
@limiter.limit("5/minute")
async def generate_proposal(
    request: Request,
    tender_id: int,
    payload: ProposalWizardRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only generate proposals for tenders you created.")

    org = _get_org(membership.organization_id, db)

    # Build tender analysis dict from stored fields
    tender_analysis = {
        "summary": tender.summary,
        "eligibility": tender.eligibility,
        "financial_requirements": tender.financial_requirements,
        "required_documents": tender.required_documents,
        "compliance_matrix": tender.compliance_matrix,
        "risk_analysis": tender.risk_analysis,
    }

    # Merge knowledge base with basic profile
    kb = _get_knowledge_base(current_user)
    if not kb.get("company_name"):
        kb["company_name"] = org.name or ""
    if not kb.get("contact_name"):
        kb["contact_name"] = current_user.contact_name or ""
    if not kb.get("phone"):
        kb["phone"] = current_user.phone or ""
    if not kb.get("address"):
        kb["address"] = current_user.address or ""

    wizard_data = payload.model_dump()
    tender_id_copy = tender.id
    organization_id = membership.organization_id
    language = payload.language

    async def generate():
        yield _sse({"type": "progress", "stage": "generating", "message": "AI is writing your personalized proposal..."})

        loop = asyncio.get_running_loop()
        queue: asyncio.Queue = asyncio.Queue()

        def produce():
            try:
                for chunk in stream_personalized_proposal(tender_analysis, kb, wizard_data, language):
                    asyncio.run_coroutine_threadsafe(queue.put(("chunk", chunk)), loop).result()
            except Exception as exc:
                asyncio.run_coroutine_threadsafe(queue.put(("error", str(exc))), loop).result()
            finally:
                asyncio.run_coroutine_threadsafe(queue.put(("done", None)), loop).result()

        thread = threading.Thread(target=produce, daemon=True)
        thread.start()

        accumulated: list[str] = []
        while True:
            kind, data = await queue.get()
            if kind == "chunk":
                accumulated.append(data)
                yield _sse({"type": "chunk", "text": data})
            elif kind == "error":
                yield _sse({"type": "error", "detail": data})
                thread.join(timeout=5)
                return
            elif kind == "done":
                break

        thread.join(timeout=5)
        yield _sse({"type": "progress", "stage": "saving", "message": "Saving personalized proposal..."})

        full_text = "".join(accumulated)

        # Save to tender
        new_db = SessionLocal()
        try:
            t = new_db.query(Tender).filter(
                Tender.id == tender_id_copy, Tender.organization_id == organization_id
            ).first()
            if t:
                t.personalized_proposal = full_text
                new_db.commit()
        finally:
            new_db.close()

        yield _sse({"type": "done", "proposal": full_text})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ---------------------------------------------------------------------------
# AI Bid Strategy Advisor
# ---------------------------------------------------------------------------

@router.post("/tenders/{tender_id}/bid-strategy")
@limiter.limit("5/minute")
async def generate_bid_strategy(
    request: Request,
    tender_id: int,
    language: str = Form("english"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only generate bid strategy for tenders you created.")

    org = _get_org(membership.organization_id, db)

    tender_analysis = {
        "summary": tender.summary,
        "eligibility": tender.eligibility,
        "financial_requirements": tender.financial_requirements,
        "required_documents": tender.required_documents,
        "compliance_matrix": tender.compliance_matrix,
        "risk_analysis": tender.risk_analysis,
        "bid_recommendation": tender.bid_recommendation,
    }

    kb = _get_knowledge_base(current_user)
    kb.setdefault("company_name", org.name or "")
    kb.setdefault("contact_name", current_user.contact_name or "")
    kb.setdefault("phone", current_user.phone or "")
    kb.setdefault("address", current_user.address or "")

    tender_id_copy = tender.id
    organization_id = membership.organization_id

    async def generate():
        yield _sse({"type": "progress", "stage": "analyzing", "message": "AI generating bid strategy, compliance analysis, and price intelligence..."})

        loop = asyncio.get_running_loop()
        queue: asyncio.Queue = asyncio.Queue()

        def produce():
            try:
                for chunk in stream_bid_strategy(tender_analysis, kb, language):
                    asyncio.run_coroutine_threadsafe(queue.put(("chunk", chunk)), loop).result()
            except Exception as exc:
                asyncio.run_coroutine_threadsafe(queue.put(("error", str(exc))), loop).result()
            finally:
                asyncio.run_coroutine_threadsafe(queue.put(("done", None)), loop).result()

        thread = threading.Thread(target=produce, daemon=True)
        thread.start()

        accumulated: list[str] = []
        while True:
            kind, data = await queue.get()
            if kind == "chunk":
                accumulated.append(data)
                yield _sse({"type": "chunk", "text": data})
            elif kind == "error":
                yield _sse({"type": "error", "detail": data})
                thread.join(timeout=5)
                return
            elif kind == "done":
                break

        thread.join(timeout=5)
        full_text = "".join(accumulated)

        new_db = SessionLocal()
        try:
            t = new_db.query(Tender).filter(
                Tender.id == tender_id_copy, Tender.organization_id == organization_id
            ).first()
            if t:
                t.bid_strategy = full_text
                new_db.commit()
        finally:
            new_db.close()

        yield _sse({"type": "done"})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )
