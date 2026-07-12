import asyncio
import hashlib
import hmac
import json
import os
import re
import secrets
import threading
from datetime import datetime, timedelta
from io import BytesIO

from fastapi import FastAPI, Request, UploadFile, File, Form, Depends, Header, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.errors import RateLimitExceeded
from slowapi.util import get_remote_address
from sqlalchemy import inspect, or_, text
from sqlalchemy.orm import Session

from database import Base, engine, SessionLocal
from models import (
    Tender, User, DiscoveredTender, Organization, OrgMembership, OrgInvite, ApprovalRequest,
    Comment, Task as TaskModel, Vendor, TenderVendorLink, Contract, Notification,
)
from hermes_client import (
    analyze_with_gemini, parse_gemini_response, stream_with_gemini,
    stream_personalized_proposal, stream_bid_strategy, validate_document,
)
from discovery import run_all_scrapers

Base.metadata.create_all(bind=engine)

# ---------------------------------------------------------------------------
# Rate limiting
# ---------------------------------------------------------------------------

def _rate_key(request: Request) -> str:
    if os.getenv("TESTING"):
        import uuid
        return str(uuid.uuid4())
    auth = request.headers.get("authorization", "")
    if auth.lower().startswith("bearer "):
        return f"user:{auth.split(' ', 1)[1].strip()}"
    return f"ip:{get_remote_address(request)}"


limiter = Limiter(key_func=_rate_key)

app = FastAPI(title="TenderOS AI Backend")
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

MAX_UPLOAD_BYTES = 15 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}
TOKEN_TTL_DAYS = 30
PLAN_LIMITS = {"free": 5, "pro": -1, "business": -1}  # -1 = unlimited

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------------------------------------------------------
# Schema migration (additive, safe for existing databases)
# ---------------------------------------------------------------------------

def ensure_schema():
    inspector = inspect(engine)

    if inspector.has_table("users"):
        existing = {c["name"] for c in inspector.get_columns("users")}
        user_cols = {
            "token_expires_at":      "TIMESTAMP",
            "plan":                  "VARCHAR(20) DEFAULT 'free'",
            "monthly_tenders_used":  "INTEGER DEFAULT 0",
            "monthly_reset_at":      "TIMESTAMP",
            "knowledge_base":        "TEXT DEFAULT '{}'",
        }
        with engine.begin() as conn:
            for col, defn in user_cols.items():
                if col not in existing:
                    conn.execute(text(f"ALTER TABLE users ADD COLUMN {col} {defn}"))

    if not inspector.has_table("discovered_tenders"):
        Base.metadata.tables["discovered_tenders"].create(bind=engine)

    if not inspector.has_table("organizations"):
        Base.metadata.tables["organizations"].create(bind=engine)

    if not inspector.has_table("org_memberships"):
        Base.metadata.tables["org_memberships"].create(bind=engine)

    if not inspector.has_table("org_invites"):
        Base.metadata.tables["org_invites"].create(bind=engine)

    if not inspector.has_table("approval_requests"):
        Base.metadata.tables["approval_requests"].create(bind=engine)

    if not inspector.has_table("comments"):
        Base.metadata.tables["comments"].create(bind=engine)

    if not inspector.has_table("tasks"):
        Base.metadata.tables["tasks"].create(bind=engine)

    if not inspector.has_table("vendors"):
        Base.metadata.tables["vendors"].create(bind=engine)

    if not inspector.has_table("tender_vendor_links"):
        Base.metadata.tables["tender_vendor_links"].create(bind=engine)

    if not inspector.has_table("contracts"):
        Base.metadata.tables["contracts"].create(bind=engine)

    if not inspector.has_table("notifications"):
        Base.metadata.tables["notifications"].create(bind=engine)

    if inspector.has_table("tenders"):
        existing = {c["name"] for c in inspector.get_columns("tenders")}
        tender_cols = {
            "user_id":               "INTEGER",
            "organization_id":       "INTEGER",
            "status":                "VARCHAR(50) DEFAULT 'completed'",
            "file_name":             "VARCHAR(255) DEFAULT ''",
            "file_type":             "VARCHAR(100) DEFAULT ''",
            "file_size":             "INTEGER DEFAULT 0",
            "deadline":              "TIMESTAMP",
            "bid_status":            "VARCHAR(50) DEFAULT 'reviewing'",
            "bid_score":             "INTEGER",
            "notes":                 "TEXT DEFAULT ''",
            "approval_status":       "VARCHAR(20) DEFAULT 'none'",
            "financial_requirements":  "TEXT",
            "bid_recommendation":     "TEXT",
            "personalized_proposal":  "TEXT",
            "bid_strategy":           "TEXT",
        }
        with engine.begin() as conn:
            for col, defn in tender_cols.items():
                if col not in existing:
                    conn.execute(text(f"ALTER TABLE tenders ADD COLUMN {col} {defn}"))

    _backfill_organizations()


def _backfill_organizations() -> None:
    """Give every user a personal Organization (idempotent, safe on every startup).

    New signups create their own org directly in the /auth/signup handler; this
    only exists to migrate users created before Organizations existed.
    """
    db = SessionLocal()
    try:
        for user in db.query(User).all():
            membership = (
                db.query(OrgMembership)
                .filter(OrgMembership.user_id == user.id)
                .order_by(OrgMembership.id.asc())
                .first()
            )
            if membership is None:
                org = Organization(
                    name=user.organization_name or f"{user.email}'s Organization",
                    plan=user.plan or "free",
                    monthly_tenders_used=user.monthly_tenders_used or 0,
                    monthly_reset_at=user.monthly_reset_at,
                )
                db.add(org)
                db.flush()
                db.add(OrgMembership(organization_id=org.id, user_id=user.id, role="owner"))
                db.commit()
                org_id = org.id
            else:
                org_id = membership.organization_id

            # Unconditional — cheap no-op once every tender has organization_id set,
            # and safe to re-run if a previous startup was interrupted mid-backfill.
            db.query(Tender).filter(
                Tender.user_id == user.id, Tender.organization_id.is_(None)
            ).update({"organization_id": org_id})
            db.commit()
    finally:
        db.close()


ensure_schema()

# ---------------------------------------------------------------------------
# Request / response models
# ---------------------------------------------------------------------------

class AuthRequest(BaseModel):
    email: str
    password: str


class ProfileUpdate(BaseModel):
    contact_name: str = ""
    phone: str = ""
    address: str = ""


class OrgUpdate(BaseModel):
    name: str


class InviteCreate(BaseModel):
    email: str
    role: str = "member"


class MemberRoleUpdate(BaseModel):
    role: str


class ApprovalDecision(BaseModel):
    decision: str
    note: str = ""


class CommentCreate(BaseModel):
    entity_type: str
    entity_id: int
    body: str


class CommentUpdate(BaseModel):
    body: str


class TaskCreate(BaseModel):
    entity_type: str | None = None
    entity_id: int | None = None
    title: str
    description: str = ""
    assignee_user_id: int | None = None
    due_date: str | None = None  # ISO-8601 date string


class TaskUpdate(BaseModel):
    title: str | None = None
    description: str | None = None
    assignee_user_id: int | None = None
    status: str | None = None
    due_date: str | None = None


class VendorCreate(BaseModel):
    name: str
    contact_name: str = ""
    email: str = ""
    phone: str = ""
    address: str = ""
    category: str = ""
    rating: int | None = None
    notes: str = ""


class VendorUpdate(BaseModel):
    name: str | None = None
    contact_name: str | None = None
    email: str | None = None
    phone: str | None = None
    address: str | None = None
    category: str | None = None
    rating: int | None = None
    notes: str | None = None


class VendorLinkCreate(BaseModel):
    vendor_id: int
    role: str = ""
    notes: str = ""


class ContractCreate(BaseModel):
    title: str
    tender_id: int | None = None
    vendor_id: int | None = None
    counterparty_name: str = ""
    contract_value: str = ""
    currency: str = "BDT"
    start_date: str | None = None
    end_date: str | None = None
    performance_security: str = ""
    notes: str = ""


class ContractUpdate(BaseModel):
    title: str | None = None
    tender_id: int | None = None
    vendor_id: int | None = None
    counterparty_name: str | None = None
    contract_value: str | None = None
    currency: str | None = None
    start_date: str | None = None
    end_date: str | None = None
    status: str | None = None
    performance_security: str | None = None
    notes: str | None = None


class ReanalyzeRequest(BaseModel):
    language: str = "english"


class TenderUpdate(BaseModel):
    bid_status: str | None = None
    notes: str | None = None
    deadline: str | None = None   # ISO-8601 date string or empty string to clear


class KnowledgeBaseUpdate(BaseModel):
    knowledge_base: dict = {}


class ProposalWizardRequest(BaseModel):
    language: str = "english"
    bid_price: str = ""
    timeline: str = ""
    warranty: str = "12 months"
    payment_terms: str = "Monthly progress payment"
    project_manager: str = ""
    methodology: str = ""

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def hash_password(password: str) -> str:
    if len(password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters.")
    salt = secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 120000)
    return f"{salt}${digest.hex()}"


def normalize_email(email: str) -> str:
    normalized = email.lower().strip()
    if "@" not in normalized or "." not in normalized.rsplit("@", 1)[-1]:
        raise HTTPException(status_code=400, detail="A valid email address is required.")
    return normalized


def verify_password(password: str, password_hash: str) -> bool:
    try:
        salt, stored = password_hash.split("$", 1)
        digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 120000).hex()
        return hmac.compare_digest(digest, stored)
    except ValueError:
        return False


def _token_expiry() -> datetime:
    return datetime.utcnow() + timedelta(days=TOKEN_TTL_DAYS)


def auth_response(user: User, org: Organization) -> dict:
    return {
        "token": user.api_token,
        "user": {
            "id": user.id,
            "email": user.email,
            "organization_name": org.name or "",
            "contact_name": user.contact_name or "",
            "phone": user.phone or "",
            "address": user.address or "",
            "plan": org.plan or "free",
        },
    }


def get_current_user(
    authorization: str | None = Header(default=None),
    db: Session = Depends(get_db),
) -> User:
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Authentication required.")

    token = authorization.split(" ", 1)[1].strip()
    user = db.query(User).filter(User.api_token == token).first()

    if not user:
        raise HTTPException(status_code=401, detail="Invalid authentication token.")

    if user.token_expires_at and user.token_expires_at < datetime.utcnow():
        raise HTTPException(status_code=401, detail="Session expired. Please log in again.")

    user.token_expires_at = _token_expiry()
    db.commit()
    return user


def _get_org(organization_id: int, db: Session) -> Organization:
    org = db.query(Organization).filter(Organization.id == organization_id).first()
    if not org:
        raise HTTPException(status_code=404, detail="Organization not found.")
    return org


def _get_user_org(user: User, db: Session) -> Organization:
    # Most-recently-joined membership wins by default (see get_current_membership).
    membership = (
        db.query(OrgMembership)
        .filter(OrgMembership.user_id == user.id, OrgMembership.status == "active")
        .order_by(OrgMembership.id.desc())
        .first()
    )
    if not membership:
        raise HTTPException(status_code=500, detail="User has no organization membership.")
    return _get_org(membership.organization_id, db)


def get_current_membership(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    x_org_id: int | None = Header(default=None, alias="X-Org-Id"),
) -> OrgMembership:
    query = db.query(OrgMembership).filter(
        OrgMembership.user_id == current_user.id,
        OrgMembership.status == "active",
    )
    if x_org_id is not None:
        query = query.filter(OrgMembership.organization_id == x_org_id)
    # Every signup auto-creates a personal org, so an invited user ends up with
    # >1 membership; default to the most recently joined one (e.g. the team
    # they just accepted an invite into) until a frontend org-switcher exists.
    membership = query.order_by(OrgMembership.id.desc()).first()
    if not membership:
        raise HTTPException(status_code=403, detail="No active organization membership.")
    return membership


def require_role(*roles: str):
    def _dep(membership: OrgMembership = Depends(get_current_membership)) -> OrgMembership:
        if membership.role not in roles:
            raise HTTPException(status_code=403, detail="You do not have permission to perform this action.")
        return membership
    return _dep


def _can_modify_tender(tender: Tender, membership: OrgMembership) -> bool:
    return membership.role in ("owner", "admin") or tender.user_id == membership.user_id


def _reset_monthly_if_needed(org: Organization, db: Session) -> None:
    now = datetime.utcnow()
    if (
        org.monthly_reset_at is None
        or org.monthly_reset_at.month != now.month
        or org.monthly_reset_at.year != now.year
    ):
        org.monthly_tenders_used = 0
        org.monthly_reset_at = now
        db.commit()


def check_usage_limit(org: Organization, db: Session) -> None:
    _reset_monthly_if_needed(org, db)
    plan = org.plan or "free"
    limit = PLAN_LIMITS.get(plan, 5)
    if limit != -1 and (org.monthly_tenders_used or 0) >= limit:
        raise HTTPException(
            status_code=402,
            detail=(
                f"Monthly limit reached ({limit} analyses/month on {plan.title()} plan). "
                "Upgrade to Pro (৳999/month) for unlimited analyses."
            ),
        )


def increment_usage(org: Organization, db: Session) -> None:
    org.monthly_tenders_used = (org.monthly_tenders_used or 0) + 1
    db.commit()


def _extract_bid_score(bid_recommendation: str | None) -> int | None:
    if not bid_recommendation:
        return None
    m = re.search(r"BID SCORE:\s*(\d+)", bid_recommendation, re.IGNORECASE)
    return max(0, min(100, int(m.group(1)))) if m else None


def extract_text_from_file(file_name: str, content: bytes) -> str:
    ext = os.path.splitext(file_name.lower())[1]
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Unsupported file type. Upload PDF, TXT, or DOCX.")

    if ext == ".pdf":
        try:
            import fitz
            pdf = fitz.open(stream=content, filetype="pdf")
            return "".join(page.get_text() for page in pdf)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Could not read PDF: {exc}")

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Could not read DOCX: {exc}")

    return content.decode("utf-8", errors="ignore")


def get_tender_response(tender: Tender) -> dict:
    return {
        "id": tender.id,
        "user_id": tender.user_id,
        "title": tender.title,
        "language": tender.language,
        "status": tender.status,
        "file_name": tender.file_name or "",
        "file_size": tender.file_size or 0,
        "deadline": tender.deadline,
        "bid_status": tender.bid_status or "reviewing",
        "bid_score": tender.bid_score,
        "notes": tender.notes or "",
        "approval_status": tender.approval_status or "none",
        "summary": tender.summary,
        "eligibility": tender.eligibility,
        "financial_requirements": tender.financial_requirements,
        "required_documents": tender.required_documents,
        "compliance_matrix": tender.compliance_matrix,
        "risk_analysis": tender.risk_analysis,
        "bid_recommendation": tender.bid_recommendation,
        "proposal_draft": tender.proposal_draft,
        "final_checklist": tender.final_checklist,
        "personalized_proposal": tender.personalized_proposal,
        "bid_strategy": tender.bid_strategy,
        "created_at": tender.created_at,
    }


def _company_profile(user: User, org: Organization) -> dict:
    return {
        "organization_name": org.name,
        "contact_name": user.contact_name,
        "email": user.email,
        "phone": user.phone,
        "address": user.address,
    }


def _get_knowledge_base(user: User) -> dict:
    try:
        return json.loads(user.knowledge_base or "{}")
    except (json.JSONDecodeError, TypeError):
        return {}


def _parse_deadline(deadline_str: str | None) -> datetime | None:
    if not deadline_str or not deadline_str.strip():
        return None
    try:
        return datetime.fromisoformat(deadline_str.strip())
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid deadline format. Use ISO-8601 (e.g. 2025-12-31).")


def _sse(data: dict) -> str:
    return f"data: {json.dumps(data, default=str)}\n\n"

# ---------------------------------------------------------------------------
# Auth routes
# ---------------------------------------------------------------------------

@app.get("/")
def home():
    return {"status": "TenderOS backend running"}


@app.post("/auth/signup")
@limiter.limit("5/minute")
def signup(request: Request, payload: AuthRequest, db: Session = Depends(get_db)):
    email = normalize_email(payload.email)
    if db.query(User).filter(User.email == email).first():
        raise HTTPException(status_code=409, detail="An account already exists for this email.")

    user = User(
        email=email,
        password_hash=hash_password(payload.password),
        api_token=secrets.token_urlsafe(32),
        token_expires_at=_token_expiry(),
        plan="free",
        monthly_tenders_used=0,
    )
    db.add(user)
    db.flush()

    org = Organization(name=f"{email}'s Organization", plan="free", monthly_tenders_used=0)
    db.add(org)
    db.flush()
    db.add(OrgMembership(organization_id=org.id, user_id=user.id, role="owner"))
    db.commit()
    db.refresh(user)
    return auth_response(user, org)


@app.post("/auth/login")
@limiter.limit("10/minute")
def login(request: Request, payload: AuthRequest, db: Session = Depends(get_db)):
    email = normalize_email(payload.email)
    user = db.query(User).filter(User.email == email).first()

    if not user or not verify_password(payload.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid email or password.")

    user.api_token = secrets.token_urlsafe(32)
    user.token_expires_at = _token_expiry()
    db.commit()
    db.refresh(user)
    return auth_response(user, _get_user_org(user, db))


@app.get("/me")
def get_me(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return auth_response(current_user, _get_user_org(current_user, db))["user"]


@app.put("/me/profile")
def update_profile(
    payload: ProfileUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    current_user.contact_name = payload.contact_name.strip()
    current_user.phone = payload.phone.strip()
    current_user.address = payload.address.strip()
    db.commit()
    db.refresh(current_user)
    return auth_response(current_user, _get_user_org(current_user, db))["user"]


@app.get("/me/knowledge-base")
def get_knowledge_base(current_user: User = Depends(get_current_user)):
    return _get_knowledge_base(current_user)


@app.put("/me/knowledge-base")
def update_knowledge_base(
    payload: KnowledgeBaseUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    current_user.knowledge_base = json.dumps(payload.knowledge_base, ensure_ascii=False)
    db.commit()
    return {"detail": "Knowledge base updated.", "knowledge_base": payload.knowledge_base}

# ---------------------------------------------------------------------------
# Organizations / Teams
# ---------------------------------------------------------------------------

VALID_ROLES = {"owner", "admin", "member"}


def _member_response(membership: OrgMembership, user: User | None) -> dict:
    return {
        "user_id": membership.user_id,
        "email": user.email if user else "",
        "contact_name": user.contact_name if user else "",
        "role": membership.role,
        "joined_at": membership.created_at,
    }


@app.get("/orgs/me")
def get_my_org(
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    org = _get_org(membership.organization_id, db)
    return {"id": org.id, "name": org.name, "plan": org.plan or "free", "role": membership.role}


@app.put("/orgs/me")
def update_my_org(
    payload: OrgUpdate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(require_role("owner")),
):
    name = payload.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="Organization name is required.")
    org = _get_org(membership.organization_id, db)
    org.name = name
    db.commit()
    return {"id": org.id, "name": org.name, "plan": org.plan or "free", "role": membership.role}


@app.get("/orgs/me/members")
def list_members(
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    memberships = (
        db.query(OrgMembership)
        .filter(
            OrgMembership.organization_id == membership.organization_id,
            OrgMembership.status == "active",
        )
        .order_by(OrgMembership.id.asc())
        .all()
    )
    users = {
        u.id: u
        for u in db.query(User).filter(User.id.in_([m.user_id for m in memberships])).all()
    }
    return [_member_response(m, users.get(m.user_id)) for m in memberships]


def _active_owner_count(organization_id: int, db: Session) -> int:
    return (
        db.query(OrgMembership)
        .filter(
            OrgMembership.organization_id == organization_id,
            OrgMembership.role == "owner",
            OrgMembership.status == "active",
        )
        .count()
    )


@app.patch("/orgs/me/members/{user_id}")
def update_member_role(
    user_id: int,
    payload: MemberRoleUpdate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(require_role("owner")),
):
    if payload.role not in VALID_ROLES:
        raise HTTPException(status_code=400, detail=f"Invalid role. Must be one of: {', '.join(sorted(VALID_ROLES))}")

    target = db.query(OrgMembership).filter(
        OrgMembership.organization_id == membership.organization_id,
        OrgMembership.user_id == user_id,
        OrgMembership.status == "active",
    ).first()
    if not target:
        raise HTTPException(status_code=404, detail="Member not found.")

    if target.role == "owner" and payload.role != "owner" and _active_owner_count(membership.organization_id, db) <= 1:
        raise HTTPException(status_code=400, detail="Cannot demote the sole remaining owner.")

    target.role = payload.role
    db.commit()
    return _member_response(target, db.query(User).filter(User.id == user_id).first())


@app.delete("/orgs/me/members/{user_id}")
def remove_member(
    user_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(require_role("owner", "admin")),
):
    if user_id == membership.user_id:
        raise HTTPException(status_code=400, detail="You cannot remove yourself from the organization.")

    target = db.query(OrgMembership).filter(
        OrgMembership.organization_id == membership.organization_id,
        OrgMembership.user_id == user_id,
        OrgMembership.status == "active",
    ).first()
    if not target:
        raise HTTPException(status_code=404, detail="Member not found.")

    if target.role == "owner" and _active_owner_count(membership.organization_id, db) <= 1:
        raise HTTPException(status_code=400, detail="Cannot remove the sole remaining owner.")
    if membership.role == "admin" and target.role in ("owner", "admin"):
        raise HTTPException(status_code=403, detail="Admins cannot remove owners or other admins.")

    target.status = "removed"
    db.commit()
    return {"detail": "Member removed."}


@app.post("/orgs/me/invites")
@limiter.limit("10/minute")
def create_invite(
    request: Request,
    payload: InviteCreate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(require_role("owner", "admin")),
):
    if payload.role not in VALID_ROLES:
        raise HTTPException(status_code=400, detail=f"Invalid role. Must be one of: {', '.join(sorted(VALID_ROLES))}")
    email = normalize_email(payload.email)

    already_member = (
        db.query(OrgMembership)
        .join(User, User.id == OrgMembership.user_id)
        .filter(
            OrgMembership.organization_id == membership.organization_id,
            OrgMembership.status == "active",
            User.email == email,
        )
        .first()
    )
    if already_member:
        raise HTTPException(status_code=409, detail="This person is already a member of your organization.")

    invite = OrgInvite(
        organization_id=membership.organization_id,
        email=email,
        role=payload.role,
        token=secrets.token_urlsafe(24),
        invited_by_user_id=membership.user_id,
        status="pending",
        expires_at=datetime.utcnow() + timedelta(days=7),
    )
    db.add(invite)
    db.commit()
    db.refresh(invite)
    return {
        "id": invite.id, "email": invite.email, "role": invite.role, "token": invite.token,
        "status": invite.status, "expires_at": invite.expires_at, "created_at": invite.created_at,
    }


@app.get("/orgs/me/invites")
def list_invites(
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(require_role("owner", "admin")),
):
    invites = (
        db.query(OrgInvite)
        .filter(OrgInvite.organization_id == membership.organization_id, OrgInvite.status == "pending")
        .order_by(OrgInvite.id.desc())
        .all()
    )
    return [
        {
            "id": i.id, "email": i.email, "role": i.role, "token": i.token,
            "status": i.status, "expires_at": i.expires_at, "created_at": i.created_at,
        }
        for i in invites
    ]


@app.delete("/orgs/me/invites/{invite_id}")
def revoke_invite(
    invite_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(require_role("owner", "admin")),
):
    invite = db.query(OrgInvite).filter(
        OrgInvite.id == invite_id, OrgInvite.organization_id == membership.organization_id
    ).first()
    if not invite:
        raise HTTPException(status_code=404, detail="Invite not found.")
    invite.status = "revoked"
    db.commit()
    return {"detail": "Invite revoked."}


@app.post("/invites/{token}/accept")
def accept_invite(
    token: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    invite = db.query(OrgInvite).filter(OrgInvite.token == token).first()
    if not invite:
        raise HTTPException(status_code=404, detail="Invite not found.")
    if invite.status != "pending":
        raise HTTPException(status_code=400, detail="This invite is no longer valid.")
    if invite.expires_at and invite.expires_at < datetime.utcnow():
        invite.status = "expired"
        db.commit()
        raise HTTPException(status_code=400, detail="This invite has expired.")
    if invite.email != current_user.email:
        raise HTTPException(status_code=403, detail="This invite was sent to a different email address.")

    existing = db.query(OrgMembership).filter(
        OrgMembership.organization_id == invite.organization_id,
        OrgMembership.user_id == current_user.id,
    ).first()
    if existing:
        existing.status = "active"
        existing.role = invite.role
    else:
        db.add(OrgMembership(
            organization_id=invite.organization_id,
            user_id=current_user.id,
            role=invite.role,
            status="active",
        ))

    invite.status = "accepted"
    _notify(
        db, invite.organization_id, invite.invited_by_user_id, "member_invited",
        title=f"{current_user.email} accepted your invite",
        entity_type=None, entity_id=None,
    )
    db.commit()
    org = _get_org(invite.organization_id, db)
    return {"detail": "Invite accepted.", "organization": {"id": org.id, "name": org.name}}

# ---------------------------------------------------------------------------
# Subscription
# ---------------------------------------------------------------------------

@app.get("/subscription")
def get_subscription(
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    org = _get_org(membership.organization_id, db)
    _reset_monthly_if_needed(org, db)
    plan = org.plan or "free"
    limit = PLAN_LIMITS.get(plan, 5)
    return {
        "plan": plan,
        "monthly_tenders_used": org.monthly_tenders_used or 0,
        "monthly_limit": limit,
        "is_unlimited": limit == -1,
    }

# ---------------------------------------------------------------------------
# Tender routes
# ---------------------------------------------------------------------------

@app.get("/tenders")
def list_tenders(
    search: str = "",
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    query = db.query(Tender).filter(Tender.organization_id == membership.organization_id)

    if search.strip():
        term = f"%{search.strip()}%"
        query = query.filter(
            or_(Tender.title.ilike(term), Tender.file_name.ilike(term))
        )

    tenders = query.order_by(Tender.id.desc()).all()
    return [
        {
            "id": t.id,
            "user_id": t.user_id,
            "title": t.title,
            "language": t.language,
            "status": t.status,
            "file_name": t.file_name or "",
            "file_size": t.file_size or 0,
            "deadline": t.deadline,
            "bid_status": t.bid_status or "reviewing",
            "bid_score": t.bid_score,
            "approval_status": t.approval_status or "none",
            "created_at": t.created_at,
            "summary": (t.summary or "")[:180],
        }
        for t in tenders
    ]


@app.get("/tenders/{tender_id}")
def get_tender(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    return get_tender_response(tender)


@app.patch("/tenders/{tender_id}")
def update_tender(
    tender_id: int,
    payload: TenderUpdate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only edit tenders you created.")

    valid_bid_statuses = {"reviewing", "submitted", "won", "lost", "no-bid"}
    if payload.bid_status is not None:
        if payload.bid_status not in valid_bid_statuses:
            raise HTTPException(status_code=400, detail=f"Invalid bid status. Must be one of: {', '.join(valid_bid_statuses)}")
        tender.bid_status = payload.bid_status

    if payload.notes is not None:
        tender.notes = payload.notes

    if payload.deadline is not None:
        tender.deadline = _parse_deadline(payload.deadline)

    db.commit()
    db.refresh(tender)
    return get_tender_response(tender)


@app.post("/tenders/analyze")
@limiter.limit("10/minute")
async def analyze_tender(
    request: Request,
    title: str = Form(...),
    language: str = Form("english"),
    deadline: str = Form(""),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    clean_title = title.strip()
    if not clean_title:
        raise HTTPException(status_code=400, detail="Tender title is required.")
    if language not in {"english", "bangla"}:
        raise HTTPException(status_code=400, detail="Unsupported language.")

    org = _get_org(membership.organization_id, db)
    check_usage_limit(org, db)

    content = await file.read()
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File is too large. Maximum upload is 15 MB.")

    tender_text = extract_text_from_file(file.filename, content)
    if len(tender_text.strip()) < 20:
        raise HTTPException(status_code=400, detail="Unable to extract readable text from this file.")

    result = analyze_with_gemini(tender_text, language, _company_profile(current_user, org))
    bid_score = _extract_bid_score(result.get("bid_recommendation"))

    tender = Tender(
        user_id=current_user.id,
        organization_id=membership.organization_id,
        title=clean_title,
        language=language,
        status="failed" if (result.get("summary") or "").startswith("Gemini API Error:") else "completed",
        file_name=file.filename,
        file_type=file.content_type or "",
        file_size=len(content),
        deadline=_parse_deadline(deadline),
        bid_score=bid_score,
        original_text=tender_text,
        summary=result.get("summary"),
        eligibility=result.get("eligibility"),
        financial_requirements=result.get("financial_requirements"),
        required_documents=result.get("required_documents"),
        compliance_matrix=result.get("compliance_matrix"),
        risk_analysis=result.get("risk_analysis"),
        bid_recommendation=result.get("bid_recommendation"),
        proposal_draft=result.get("proposal_draft"),
        final_checklist=result.get("final_checklist"),
    )
    db.add(tender)
    db.commit()
    db.refresh(tender)
    increment_usage(org, db)
    return get_tender_response(tender)


@app.post("/tenders/analyze-stream")
@limiter.limit("10/minute")
async def analyze_tender_stream(
    request: Request,
    title: str = Form(...),
    language: str = Form("english"),
    deadline: str = Form(""),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    clean_title = title.strip()
    if not clean_title:
        raise HTTPException(status_code=400, detail="Tender title is required.")
    if language not in {"english", "bangla"}:
        raise HTTPException(status_code=400, detail="Unsupported language.")

    org = _get_org(membership.organization_id, db)
    check_usage_limit(org, db)

    content = await file.read()
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File is too large. Maximum upload is 15 MB.")

    tender_text = extract_text_from_file(file.filename, content)
    if len(tender_text.strip()) < 20:
        raise HTTPException(status_code=400, detail="Unable to extract readable text from this file.")

    file_name = file.filename
    file_type = file.content_type or ""
    file_size = len(content)
    user_id = current_user.id
    organization_id = membership.organization_id
    profile = _company_profile(current_user, org)
    parsed_deadline = _parse_deadline(deadline)

    async def generate():
        yield _sse({"type": "progress", "stage": "analyzing", "message": "AI is analyzing your tender..."})

        loop = asyncio.get_running_loop()
        queue: asyncio.Queue = asyncio.Queue()

        def produce():
            try:
                for chunk in stream_with_gemini(tender_text, language, profile):
                    asyncio.run_coroutine_threadsafe(queue.put(("chunk", chunk)), loop).result()
            except Exception as exc:
                asyncio.run_coroutine_threadsafe(queue.put(("error", str(exc))), loop).result()
            finally:
                asyncio.run_coroutine_threadsafe(queue.put(("done", None)), loop).result()

        thread = threading.Thread(target=produce, daemon=True)
        thread.start()

        accumulated: list[str] = []
        while True:
            kind, data = await queue.get()
            if kind == "chunk":
                accumulated.append(data)
                yield _sse({"type": "chunk", "text": data})
            elif kind == "error":
                yield _sse({"type": "error", "detail": data})
                thread.join(timeout=5)
                return
            elif kind == "done":
                break

        thread.join(timeout=5)
        yield _sse({"type": "progress", "stage": "saving", "message": "Saving results..."})

        full_text = "".join(accumulated)
        result = parse_gemini_response(full_text)
        bid_score = _extract_bid_score(result.get("bid_recommendation"))
        status = "failed" if (result.get("summary") or "").startswith("Gemini API Error:") else "completed"

        tender = Tender(
            user_id=user_id,
            organization_id=organization_id,
            title=clean_title,
            language=language,
            status=status,
            file_name=file_name,
            file_type=file_type,
            file_size=file_size,
            deadline=parsed_deadline,
            bid_score=bid_score,
            original_text=tender_text,
            summary=result.get("summary"),
            eligibility=result.get("eligibility"),
            financial_requirements=result.get("financial_requirements"),
            required_documents=result.get("required_documents"),
            compliance_matrix=result.get("compliance_matrix"),
            risk_analysis=result.get("risk_analysis"),
            bid_recommendation=result.get("bid_recommendation"),
            proposal_draft=result.get("proposal_draft"),
            final_checklist=result.get("final_checklist"),
        )
        db.add(tender)
        db.commit()
        db.refresh(tender)
        increment_usage(org, db)

        yield _sse({"type": "done", "tender": get_tender_response(tender)})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/tenders/{tender_id}/reanalyze")
@limiter.limit("10/minute")
def reanalyze_tender(
    request: Request,
    tender_id: int,
    payload: ReanalyzeRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only re-analyze tenders you created.")
    if payload.language not in {"english", "bangla"}:
        raise HTTPException(status_code=400, detail="Unsupported language.")
    if not tender.original_text:
        raise HTTPException(status_code=400, detail="Original document text not available for re-analysis.")

    org = _get_org(membership.organization_id, db)
    check_usage_limit(org, db)

    result = analyze_with_gemini(tender.original_text, payload.language, _company_profile(current_user, org))
    bid_score = _extract_bid_score(result.get("bid_recommendation"))

    tender.language = payload.language
    tender.status = "failed" if (result.get("summary") or "").startswith("Gemini API Error:") else "completed"
    tender.summary = result.get("summary")
    tender.eligibility = result.get("eligibility")
    tender.financial_requirements = result.get("financial_requirements")
    tender.required_documents = result.get("required_documents")
    tender.compliance_matrix = result.get("compliance_matrix")
    tender.risk_analysis = result.get("risk_analysis")
    tender.bid_recommendation = result.get("bid_recommendation")
    tender.proposal_draft = result.get("proposal_draft")
    tender.final_checklist = result.get("final_checklist")
    tender.bid_score = bid_score

    db.commit()
    db.refresh(tender)
    increment_usage(org, db)
    return get_tender_response(tender)


@app.delete("/tenders/{tender_id}")
def delete_tender(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only delete tenders you created.")
    db.delete(tender)
    db.commit()
    return {"detail": "Tender deleted."}


@app.get("/tenders/{tender_id}/export-docx")
def export_docx(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")

    from docx import Document

    doc = Document()
    doc.add_heading("TenderOS AI — Tender Analysis Report", 0)
    doc.add_heading(tender.title, level=1)
    doc.add_paragraph(f"Output Language: {tender.language.title()}")
    if tender.deadline:
        doc.add_paragraph(f"Submission Deadline: {tender.deadline.strftime('%d %B %Y')}")
    if tender.bid_score is not None:
        doc.add_paragraph(f"AI Bid Score: {tender.bid_score}/100")
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%d %B %Y')}")

    sections = [
        ("Executive Summary",           tender.summary),
        ("Eligibility Criteria",        tender.eligibility),
        ("Financial Requirements",      tender.financial_requirements),
        ("Required Documents",          tender.required_documents),
        ("Compliance Matrix",           tender.compliance_matrix),
        ("Risk Analysis",               tender.risk_analysis),
        ("Bid Recommendation",          tender.bid_recommendation),
        ("Tender Submission Draft",     tender.proposal_draft),
        ("Final Submission Checklist",  tender.final_checklist),
    ]
    for heading, content in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(content or "Not available")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=tender_{tender.id}_report.docx"},
    )


@app.get("/tenders/{tender_id}/export-pdf")
def export_pdf(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")

    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(120, 120, 120)
            self.cell(0, 6, "TenderOS AI — Confidential Analysis Report", align="R")
            self.ln(8)

        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", size=8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 6, f"Page {self.page_no()}", align="C")

    pdf = PDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # Title block
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(30, 60, 120)
    pdf.multi_cell(0, 10, "TenderOS AI", align="C")
    pdf.set_font("Helvetica", size=11)
    pdf.set_text_color(60, 60, 60)
    pdf.multi_cell(0, 7, "Bangladesh Procurement Analysis Report", align="C")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(20, 20, 20)

    def safe(txt: str) -> str:
        return (txt or "").encode("latin-1", errors="replace").decode("latin-1")

    pdf.multi_cell(0, 8, safe(tender.title), align="C")
    pdf.ln(2)

    pdf.set_font("Helvetica", size=9)
    pdf.set_text_color(100, 100, 100)
    meta_parts = [f"Language: {tender.language.title()}"]
    if tender.deadline:
        meta_parts.append(f"Deadline: {tender.deadline.strftime('%d %b %Y')}")
    if tender.bid_score is not None:
        meta_parts.append(f"Bid Score: {tender.bid_score}/100")
    pdf.cell(0, 6, "  |  ".join(meta_parts), align="C")
    pdf.ln(10)

    pdf.set_draw_color(200, 210, 230)
    pdf.set_line_width(0.4)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(8)
    pdf.set_text_color(0, 0, 0)

    report_sections = [
        ("Executive Summary",           tender.summary),
        ("Eligibility Criteria",        tender.eligibility),
        ("Financial Requirements",      tender.financial_requirements),
        ("Required Documents",          tender.required_documents),
        ("Compliance Matrix",           tender.compliance_matrix),
        ("Risk Analysis",               tender.risk_analysis),
        ("Bid Recommendation",          tender.bid_recommendation),
        ("Tender Submission Draft",     tender.proposal_draft),
        ("Final Submission Checklist",  tender.final_checklist),
    ]

    for heading, content in report_sections:
        if not content:
            continue
        # Section heading
        pdf.set_fill_color(235, 241, 255)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(30, 60, 120)
        pdf.multi_cell(0, 8, heading, fill=True)
        pdf.ln(2)
        # Section body
        pdf.set_font("Helvetica", size=9)
        pdf.set_text_color(30, 30, 30)
        try:
            pdf.multi_cell(0, 5.5, safe(content))
        except Exception:
            pdf.multi_cell(0, 5.5, "[Content not renderable in PDF. Use DOCX export for Bengali text.]")
        pdf.ln(6)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=tender_{tender.id}_report.pdf"},
    )


# ---------------------------------------------------------------------------
# Notifications (event-driven inserts — see also the Calendar & Notifications
# section further down for the read/list endpoints and computed-on-read reminders)
# ---------------------------------------------------------------------------

def _notify(
    db: Session, organization_id: int, user_id: int, type: str,
    title: str, message: str = "", urgency: str = "info",
    entity_type: str | None = None, entity_id: int | None = None,
) -> None:
    db.add(Notification(
        organization_id=organization_id, user_id=user_id, type=type,
        entity_type=entity_type, entity_id=entity_id,
        title=title, message=message, urgency=urgency,
    ))


def _notify_admins(
    db: Session, organization_id: int, exclude_user_id: int, type: str,
    title: str, message: str = "", urgency: str = "info",
    entity_type: str | None = None, entity_id: int | None = None,
) -> None:
    admins = db.query(OrgMembership).filter(
        OrgMembership.organization_id == organization_id,
        OrgMembership.status == "active",
        OrgMembership.role.in_(("owner", "admin")),
        OrgMembership.user_id != exclude_user_id,
    ).all()
    for m in admins:
        _notify(db, organization_id, m.user_id, type, title, message, urgency, entity_type, entity_id)


# ---------------------------------------------------------------------------
# Approval Workflow
# ---------------------------------------------------------------------------

VALID_APPROVAL_DECISIONS = {"approved", "rejected"}


def _approval_response(a: ApprovalRequest) -> dict:
    return {
        "id": a.id,
        "tender_id": a.tender_id,
        "requested_by_user_id": a.requested_by_user_id,
        "status": a.status,
        "reviewer_user_id": a.reviewer_user_id,
        "reviewer_note": a.reviewer_note or "",
        "requested_at": a.requested_at,
        "reviewed_at": a.reviewed_at,
    }


def _get_org_tender(tender_id: int, membership: OrgMembership, db: Session) -> Tender:
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    return tender


@app.post("/tenders/{tender_id}/approval/request")
def request_approval(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = _get_org_tender(tender_id, membership, db)
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only request approval for tenders you created.")
    if tender.approval_status == "pending":
        raise HTTPException(status_code=400, detail="This tender already has a pending approval request.")

    approval = ApprovalRequest(
        organization_id=membership.organization_id,
        tender_id=tender.id,
        requested_by_user_id=membership.user_id,
        status="pending",
    )
    db.add(approval)
    tender.approval_status = "pending"
    _notify_admins(
        db, membership.organization_id, membership.user_id, "approval_requested",
        title=f'Approval requested: "{tender.title}"',
        urgency="warning", entity_type="tender", entity_id=tender.id,
    )
    db.commit()
    db.refresh(approval)
    return _approval_response(approval)


@app.post("/tenders/{tender_id}/approval/cancel")
def cancel_approval(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = _get_org_tender(tender_id, membership, db)
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only cancel approval requests for tenders you created.")

    pending = (
        db.query(ApprovalRequest)
        .filter(ApprovalRequest.tender_id == tender.id, ApprovalRequest.status == "pending")
        .order_by(ApprovalRequest.id.desc())
        .first()
    )
    if not pending:
        raise HTTPException(status_code=400, detail="This tender has no pending approval request.")

    pending.status = "cancelled"
    pending.reviewed_at = datetime.utcnow()
    tender.approval_status = "none"
    db.commit()
    return {"detail": "Approval request cancelled."}


@app.post("/tenders/{tender_id}/approval/decide")
def decide_approval(
    tender_id: int,
    payload: ApprovalDecision,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(require_role("owner", "admin")),
):
    if payload.decision not in VALID_APPROVAL_DECISIONS:
        raise HTTPException(status_code=400, detail=f"Invalid decision. Must be one of: {', '.join(sorted(VALID_APPROVAL_DECISIONS))}")

    tender = _get_org_tender(tender_id, membership, db)
    pending = (
        db.query(ApprovalRequest)
        .filter(ApprovalRequest.tender_id == tender.id, ApprovalRequest.status == "pending")
        .order_by(ApprovalRequest.id.desc())
        .first()
    )
    if not pending:
        raise HTTPException(status_code=400, detail="This tender has no pending approval request.")

    pending.status = payload.decision
    pending.reviewer_user_id = membership.user_id
    pending.reviewer_note = payload.note.strip()
    pending.reviewed_at = datetime.utcnow()
    tender.approval_status = payload.decision
    if pending.requested_by_user_id != membership.user_id:
        _notify(
            db, membership.organization_id, pending.requested_by_user_id, "approval_decided",
            title=f'Tender "{tender.title}" was {payload.decision}',
            message=pending.reviewer_note,
            urgency="critical" if payload.decision == "rejected" else "info",
            entity_type="tender", entity_id=tender.id,
        )
    db.commit()
    return _approval_response(pending)


@app.get("/tenders/{tender_id}/approval/history")
def get_approval_history(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = _get_org_tender(tender_id, membership, db)
    history = (
        db.query(ApprovalRequest)
        .filter(ApprovalRequest.tender_id == tender.id)
        .order_by(ApprovalRequest.id.desc())
        .all()
    )
    return [_approval_response(a) for a in history]


@app.get("/approvals/pending")
def list_pending_approvals(
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(require_role("owner", "admin")),
):
    pending = (
        db.query(ApprovalRequest)
        .filter(
            ApprovalRequest.organization_id == membership.organization_id,
            ApprovalRequest.status == "pending",
        )
        .order_by(ApprovalRequest.id.asc())
        .all()
    )
    tenders = {
        t.id: t
        for t in db.query(Tender).filter(Tender.id.in_([a.tender_id for a in pending])).all()
    }
    return [
        {
            **_approval_response(a),
            "tender_title": tenders[a.tender_id].title if a.tender_id in tenders else "",
        }
        for a in pending
    ]


# ---------------------------------------------------------------------------
# Comments & Tasks
# ---------------------------------------------------------------------------

VALID_ENTITY_TYPES = {"tender", "vendor", "contract"}
VALID_TASK_STATUSES = {"open", "in_progress", "done", "cancelled"}


def _validate_entity(entity_type: str, entity_id: int, organization_id: int, db: Session) -> None:
    if entity_type not in VALID_ENTITY_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported entity_type. Must be one of: {', '.join(sorted(VALID_ENTITY_TYPES))}",
        )
    if entity_type == "tender":
        exists = db.query(Tender).filter(
            Tender.id == entity_id, Tender.organization_id == organization_id
        ).first()
        if not exists:
            raise HTTPException(status_code=404, detail="Tender not found.")
    elif entity_type == "vendor":
        exists = db.query(Vendor).filter(
            Vendor.id == entity_id, Vendor.organization_id == organization_id
        ).first()
        if not exists:
            raise HTTPException(status_code=404, detail="Vendor not found.")
    elif entity_type == "contract":
        exists = db.query(Contract).filter(
            Contract.id == entity_id, Contract.organization_id == organization_id
        ).first()
        if not exists:
            raise HTTPException(status_code=404, detail="Contract not found.")


def _users_by_id(user_ids: list[int | None], db: Session) -> dict[int, User]:
    ids = [uid for uid in user_ids if uid is not None]
    if not ids:
        return {}
    return {u.id: u for u in db.query(User).filter(User.id.in_(ids)).all()}


def _comment_response(c: Comment, users: dict[int, User]) -> dict:
    author = users.get(c.author_user_id)
    return {
        "id": c.id,
        "entity_type": c.entity_type,
        "entity_id": c.entity_id,
        "author_user_id": c.author_user_id,
        "author_email": author.email if author else "",
        "body": c.body,
        "created_at": c.created_at,
        "updated_at": c.updated_at,
    }


def _task_response(t: TaskModel, users: dict[int, User]) -> dict:
    assignee = users.get(t.assignee_user_id) if t.assignee_user_id else None
    creator = users.get(t.created_by_user_id)
    return {
        "id": t.id,
        "entity_type": t.entity_type,
        "entity_id": t.entity_id,
        "title": t.title,
        "description": t.description or "",
        "assignee_user_id": t.assignee_user_id,
        "assignee_email": assignee.email if assignee else "",
        "created_by_user_id": t.created_by_user_id,
        "created_by_email": creator.email if creator else "",
        "status": t.status,
        "due_date": t.due_date,
        "created_at": t.created_at,
        "updated_at": t.updated_at,
    }


@app.get("/comments")
def list_comments(
    entity_type: str,
    entity_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    _validate_entity(entity_type, entity_id, membership.organization_id, db)
    comments = (
        db.query(Comment)
        .filter(
            Comment.organization_id == membership.organization_id,
            Comment.entity_type == entity_type,
            Comment.entity_id == entity_id,
        )
        .order_by(Comment.id.asc())
        .all()
    )
    users = _users_by_id([c.author_user_id for c in comments], db)
    return [_comment_response(c, users) for c in comments]


def _entity_owner_user_id(entity_type: str, entity_id: int, db: Session) -> int | None:
    if entity_type == "tender":
        row = db.query(Tender).filter(Tender.id == entity_id).first()
        return row.user_id if row else None
    if entity_type == "vendor":
        row = db.query(Vendor).filter(Vendor.id == entity_id).first()
        return row.created_by_user_id if row else None
    if entity_type == "contract":
        row = db.query(Contract).filter(Contract.id == entity_id).first()
        return row.created_by_user_id if row else None
    return None


def _entity_title(entity_type: str, entity_id: int, db: Session) -> str:
    if entity_type == "tender":
        row = db.query(Tender).filter(Tender.id == entity_id).first()
        return row.title if row else "tender"
    if entity_type == "vendor":
        row = db.query(Vendor).filter(Vendor.id == entity_id).first()
        return row.name if row else "vendor"
    if entity_type == "contract":
        row = db.query(Contract).filter(Contract.id == entity_id).first()
        return row.title if row else "contract"
    return entity_type


@app.post("/comments")
def create_comment(
    payload: CommentCreate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    body = payload.body.strip()
    if not body:
        raise HTTPException(status_code=400, detail="Comment body is required.")
    _validate_entity(payload.entity_type, payload.entity_id, membership.organization_id, db)

    comment = Comment(
        organization_id=membership.organization_id,
        entity_type=payload.entity_type,
        entity_id=payload.entity_id,
        author_user_id=membership.user_id,
        body=body,
    )
    db.add(comment)

    owner_id = _entity_owner_user_id(payload.entity_type, payload.entity_id, db)
    if owner_id and owner_id != membership.user_id:
        _notify(
            db, membership.organization_id, owner_id, "comment_added",
            title=f'New comment on "{_entity_title(payload.entity_type, payload.entity_id, db)}"',
            message=body[:200],
            entity_type=payload.entity_type, entity_id=payload.entity_id,
        )

    db.commit()
    db.refresh(comment)
    return _comment_response(comment, _users_by_id([comment.author_user_id], db))


@app.patch("/comments/{comment_id}")
def update_comment(
    comment_id: int,
    payload: CommentUpdate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    comment = db.query(Comment).filter(
        Comment.id == comment_id, Comment.organization_id == membership.organization_id
    ).first()
    if not comment:
        raise HTTPException(status_code=404, detail="Comment not found.")
    if comment.author_user_id != membership.user_id:
        raise HTTPException(status_code=403, detail="You can only edit your own comments.")

    body = payload.body.strip()
    if not body:
        raise HTTPException(status_code=400, detail="Comment body is required.")
    comment.body = body
    comment.updated_at = datetime.utcnow()
    db.commit()
    return _comment_response(comment, _users_by_id([comment.author_user_id], db))


@app.delete("/comments/{comment_id}")
def delete_comment(
    comment_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    comment = db.query(Comment).filter(
        Comment.id == comment_id, Comment.organization_id == membership.organization_id
    ).first()
    if not comment:
        raise HTTPException(status_code=404, detail="Comment not found.")
    if comment.author_user_id != membership.user_id and membership.role not in ("owner", "admin"):
        raise HTTPException(status_code=403, detail="You can only delete your own comments.")

    db.delete(comment)
    db.commit()
    return {"detail": "Comment deleted."}


def _can_modify_task(task: TaskModel, membership: OrgMembership) -> bool:
    return (
        membership.role in ("owner", "admin")
        or task.created_by_user_id == membership.user_id
        or task.assignee_user_id == membership.user_id
    )


@app.get("/tasks")
def list_tasks(
    entity_type: str | None = None,
    entity_id: int | None = None,
    assignee_user_id: int | None = None,
    status: str | None = None,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    query = db.query(TaskModel).filter(TaskModel.organization_id == membership.organization_id)
    if entity_type is not None:
        query = query.filter(TaskModel.entity_type == entity_type)
    if entity_id is not None:
        query = query.filter(TaskModel.entity_id == entity_id)
    if assignee_user_id is not None:
        query = query.filter(TaskModel.assignee_user_id == assignee_user_id)
    if status is not None:
        query = query.filter(TaskModel.status == status)

    tasks = query.order_by(TaskModel.id.desc()).all()
    users = _users_by_id(
        [t.assignee_user_id for t in tasks if t.assignee_user_id] + [t.created_by_user_id for t in tasks], db
    )
    return [_task_response(t, users) for t in tasks]


@app.get("/tasks/mine")
def list_my_tasks(
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tasks = (
        db.query(TaskModel)
        .filter(
            TaskModel.organization_id == membership.organization_id,
            TaskModel.assignee_user_id == membership.user_id,
        )
        .order_by(TaskModel.id.desc())
        .all()
    )
    users = _users_by_id(
        [t.assignee_user_id for t in tasks] + [t.created_by_user_id for t in tasks], db
    )
    return [_task_response(t, users) for t in tasks]


@app.post("/tasks")
def create_task(
    payload: TaskCreate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    title = payload.title.strip()
    if not title:
        raise HTTPException(status_code=400, detail="Task title is required.")
    if payload.entity_type is not None:
        if payload.entity_id is None:
            raise HTTPException(status_code=400, detail="entity_id is required when entity_type is set.")
        _validate_entity(payload.entity_type, payload.entity_id, membership.organization_id, db)

    task = TaskModel(
        organization_id=membership.organization_id,
        entity_type=payload.entity_type,
        entity_id=payload.entity_id,
        title=title,
        description=payload.description,
        assignee_user_id=payload.assignee_user_id,
        created_by_user_id=membership.user_id,
        status="open",
        due_date=_parse_deadline(payload.due_date),
    )
    db.add(task)

    if payload.assignee_user_id and payload.assignee_user_id != membership.user_id:
        _notify(
            db, membership.organization_id, payload.assignee_user_id, "task_assigned",
            title=f'You were assigned: "{title}"',
            entity_type=payload.entity_type, entity_id=payload.entity_id,
        )

    db.commit()
    db.refresh(task)
    return _task_response(task, _users_by_id([task.assignee_user_id, task.created_by_user_id], db))


@app.patch("/tasks/{task_id}")
def update_task(
    task_id: int,
    payload: TaskUpdate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    task = db.query(TaskModel).filter(
        TaskModel.id == task_id, TaskModel.organization_id == membership.organization_id
    ).first()
    if not task:
        raise HTTPException(status_code=404, detail="Task not found.")
    if not _can_modify_task(task, membership):
        raise HTTPException(status_code=403, detail="You can only edit tasks you created or are assigned to.")

    if payload.title is not None:
        title = payload.title.strip()
        if not title:
            raise HTTPException(status_code=400, detail="Task title is required.")
        task.title = title
    if payload.description is not None:
        task.description = payload.description
    if payload.assignee_user_id is not None:
        reassigned = payload.assignee_user_id != task.assignee_user_id
        task.assignee_user_id = payload.assignee_user_id
        if reassigned and payload.assignee_user_id != membership.user_id:
            _notify(
                db, membership.organization_id, payload.assignee_user_id, "task_assigned",
                title=f'You were assigned: "{task.title}"',
                entity_type=task.entity_type, entity_id=task.entity_id,
            )
    if payload.status is not None:
        if payload.status not in VALID_TASK_STATUSES:
            raise HTTPException(status_code=400, detail=f"Invalid status. Must be one of: {', '.join(sorted(VALID_TASK_STATUSES))}")
        task.status = payload.status
    if payload.due_date is not None:
        task.due_date = _parse_deadline(payload.due_date)

    task.updated_at = datetime.utcnow()
    db.commit()
    return _task_response(task, _users_by_id([task.assignee_user_id, task.created_by_user_id], db))


@app.delete("/tasks/{task_id}")
def delete_task(
    task_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    task = db.query(TaskModel).filter(
        TaskModel.id == task_id, TaskModel.organization_id == membership.organization_id
    ).first()
    if not task:
        raise HTTPException(status_code=404, detail="Task not found.")
    if task.created_by_user_id != membership.user_id and membership.role not in ("owner", "admin"):
        raise HTTPException(status_code=403, detail="You can only delete tasks you created.")

    db.delete(task)
    db.commit()
    return {"detail": "Task deleted."}


# ---------------------------------------------------------------------------
# Vendor Management
# ---------------------------------------------------------------------------

def _can_modify_vendor(vendor: Vendor, membership: OrgMembership) -> bool:
    return membership.role in ("owner", "admin") or vendor.created_by_user_id == membership.user_id


def _vendor_response(v: Vendor) -> dict:
    return {
        "id": v.id,
        "name": v.name,
        "contact_name": v.contact_name or "",
        "email": v.email or "",
        "phone": v.phone or "",
        "address": v.address or "",
        "category": v.category or "",
        "rating": v.rating,
        "notes": v.notes or "",
        "created_by_user_id": v.created_by_user_id,
        "created_at": v.created_at,
        "updated_at": v.updated_at,
    }


def _get_org_vendor(vendor_id: int, membership: OrgMembership, db: Session) -> Vendor:
    vendor = db.query(Vendor).filter(
        Vendor.id == vendor_id, Vendor.organization_id == membership.organization_id
    ).first()
    if not vendor:
        raise HTTPException(status_code=404, detail="Vendor not found.")
    return vendor


@app.get("/vendors")
def list_vendors(
    search: str = "",
    category: str = "",
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    query = db.query(Vendor).filter(Vendor.organization_id == membership.organization_id)
    if search.strip():
        term = f"%{search.strip()}%"
        query = query.filter(or_(Vendor.name.ilike(term), Vendor.contact_name.ilike(term)))
    if category.strip():
        query = query.filter(Vendor.category == category.strip())
    vendors = query.order_by(Vendor.name.asc()).all()
    return [_vendor_response(v) for v in vendors]


@app.post("/vendors")
def create_vendor(
    payload: VendorCreate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    name = payload.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="Vendor name is required.")

    vendor = Vendor(
        organization_id=membership.organization_id,
        name=name,
        contact_name=payload.contact_name.strip(),
        email=payload.email.strip(),
        phone=payload.phone.strip(),
        address=payload.address.strip(),
        category=payload.category.strip(),
        rating=payload.rating,
        notes=payload.notes,
        created_by_user_id=membership.user_id,
    )
    db.add(vendor)
    db.commit()
    db.refresh(vendor)
    return _vendor_response(vendor)


@app.get("/vendors/{vendor_id}")
def get_vendor(
    vendor_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    return _vendor_response(_get_org_vendor(vendor_id, membership, db))


@app.patch("/vendors/{vendor_id}")
def update_vendor(
    vendor_id: int,
    payload: VendorUpdate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    vendor = _get_org_vendor(vendor_id, membership, db)
    if not _can_modify_vendor(vendor, membership):
        raise HTTPException(status_code=403, detail="You can only edit vendors you created.")

    if payload.name is not None:
        name = payload.name.strip()
        if not name:
            raise HTTPException(status_code=400, detail="Vendor name is required.")
        vendor.name = name
    if payload.contact_name is not None:
        vendor.contact_name = payload.contact_name.strip()
    if payload.email is not None:
        vendor.email = payload.email.strip()
    if payload.phone is not None:
        vendor.phone = payload.phone.strip()
    if payload.address is not None:
        vendor.address = payload.address.strip()
    if payload.category is not None:
        vendor.category = payload.category.strip()
    if payload.rating is not None:
        vendor.rating = payload.rating
    if payload.notes is not None:
        vendor.notes = payload.notes

    vendor.updated_at = datetime.utcnow()
    db.commit()
    return _vendor_response(vendor)


@app.delete("/vendors/{vendor_id}")
def delete_vendor(
    vendor_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    vendor = _get_org_vendor(vendor_id, membership, db)
    if not _can_modify_vendor(vendor, membership):
        raise HTTPException(status_code=403, detail="You can only delete vendors you created.")

    db.delete(vendor)
    db.commit()
    return {"detail": "Vendor deleted."}


@app.get("/tenders/{tender_id}/vendors")
def list_tender_vendors(
    tender_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = _get_org_tender(tender_id, membership, db)
    links = db.query(TenderVendorLink).filter(TenderVendorLink.tender_id == tender.id).all()
    vendors = {v.id: v for v in db.query(Vendor).filter(Vendor.id.in_([l.vendor_id for l in links])).all()}
    return [
        {
            "link_id": link.id,
            "role": link.role or "",
            "notes": link.notes or "",
            "created_at": link.created_at,
            "vendor": _vendor_response(vendors[link.vendor_id]) if link.vendor_id in vendors else None,
        }
        for link in links
    ]


@app.post("/tenders/{tender_id}/vendors")
def link_tender_vendor(
    tender_id: int,
    payload: VendorLinkCreate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = _get_org_tender(tender_id, membership, db)
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only link vendors to tenders you created.")
    _get_org_vendor(payload.vendor_id, membership, db)

    existing = db.query(TenderVendorLink).filter(
        TenderVendorLink.tender_id == tender.id, TenderVendorLink.vendor_id == payload.vendor_id
    ).first()
    if existing:
        raise HTTPException(status_code=409, detail="This vendor is already linked to this tender.")

    link = TenderVendorLink(
        tender_id=tender.id, vendor_id=payload.vendor_id, role=payload.role, notes=payload.notes,
    )
    db.add(link)
    db.commit()
    db.refresh(link)
    vendor = _get_org_vendor(payload.vendor_id, membership, db)
    return {
        "link_id": link.id, "role": link.role or "", "notes": link.notes or "",
        "created_at": link.created_at, "vendor": _vendor_response(vendor),
    }


@app.delete("/tenders/{tender_id}/vendors/{vendor_id}")
def unlink_tender_vendor(
    tender_id: int,
    vendor_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = _get_org_tender(tender_id, membership, db)
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only unlink vendors from tenders you created.")

    link = db.query(TenderVendorLink).filter(
        TenderVendorLink.tender_id == tender.id, TenderVendorLink.vendor_id == vendor_id
    ).first()
    if not link:
        raise HTTPException(status_code=404, detail="This vendor is not linked to this tender.")

    db.delete(link)
    db.commit()
    return {"detail": "Vendor unlinked."}


# ---------------------------------------------------------------------------
# Contract Tracking
# ---------------------------------------------------------------------------

VALID_CONTRACT_STATUSES = {"draft", "active", "completed", "terminated"}


def _can_modify_contract(contract: Contract, membership: OrgMembership) -> bool:
    return membership.role in ("owner", "admin") or contract.created_by_user_id == membership.user_id


def _contract_response(c: Contract) -> dict:
    return {
        "id": c.id,
        "tender_id": c.tender_id,
        "vendor_id": c.vendor_id,
        "title": c.title,
        "counterparty_name": c.counterparty_name or "",
        "contract_value": c.contract_value or "",
        "currency": c.currency or "BDT",
        "start_date": c.start_date,
        "end_date": c.end_date,
        "status": c.status or "draft",
        "performance_security": c.performance_security or "",
        "notes": c.notes or "",
        "created_by_user_id": c.created_by_user_id,
        "created_at": c.created_at,
        "updated_at": c.updated_at,
    }


def _get_org_contract(contract_id: int, membership: OrgMembership, db: Session) -> Contract:
    contract = db.query(Contract).filter(
        Contract.id == contract_id, Contract.organization_id == membership.organization_id
    ).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found.")
    return contract


@app.get("/contracts")
def list_contracts(
    status: str | None = None,
    vendor_id: int | None = None,
    tender_id: int | None = None,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    query = db.query(Contract).filter(Contract.organization_id == membership.organization_id)
    if status is not None:
        query = query.filter(Contract.status == status)
    if vendor_id is not None:
        query = query.filter(Contract.vendor_id == vendor_id)
    if tender_id is not None:
        query = query.filter(Contract.tender_id == tender_id)
    contracts = query.order_by(Contract.id.desc()).all()
    return [_contract_response(c) for c in contracts]


@app.post("/contracts")
def create_contract(
    payload: ContractCreate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    title = payload.title.strip()
    if not title:
        raise HTTPException(status_code=400, detail="Contract title is required.")
    if payload.tender_id is not None:
        _get_org_tender(payload.tender_id, membership, db)
    if payload.vendor_id is not None:
        _get_org_vendor(payload.vendor_id, membership, db)

    contract = Contract(
        organization_id=membership.organization_id,
        tender_id=payload.tender_id,
        vendor_id=payload.vendor_id,
        title=title,
        counterparty_name=payload.counterparty_name.strip(),
        contract_value=payload.contract_value.strip(),
        currency=payload.currency.strip() or "BDT",
        start_date=_parse_deadline(payload.start_date),
        end_date=_parse_deadline(payload.end_date),
        status="draft",
        performance_security=payload.performance_security.strip(),
        notes=payload.notes,
        created_by_user_id=membership.user_id,
    )
    db.add(contract)
    db.commit()
    db.refresh(contract)
    return _contract_response(contract)


@app.get("/contracts/{contract_id}")
def get_contract(
    contract_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    return _contract_response(_get_org_contract(contract_id, membership, db))


@app.patch("/contracts/{contract_id}")
def update_contract(
    contract_id: int,
    payload: ContractUpdate,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    contract = _get_org_contract(contract_id, membership, db)
    if not _can_modify_contract(contract, membership):
        raise HTTPException(status_code=403, detail="You can only edit contracts you created.")

    if payload.title is not None:
        title = payload.title.strip()
        if not title:
            raise HTTPException(status_code=400, detail="Contract title is required.")
        contract.title = title
    if payload.tender_id is not None:
        _get_org_tender(payload.tender_id, membership, db)
        contract.tender_id = payload.tender_id
    if payload.vendor_id is not None:
        _get_org_vendor(payload.vendor_id, membership, db)
        contract.vendor_id = payload.vendor_id
    if payload.counterparty_name is not None:
        contract.counterparty_name = payload.counterparty_name.strip()
    if payload.contract_value is not None:
        contract.contract_value = payload.contract_value.strip()
    if payload.currency is not None:
        contract.currency = payload.currency.strip() or "BDT"
    if payload.start_date is not None:
        contract.start_date = _parse_deadline(payload.start_date)
    if payload.end_date is not None:
        contract.end_date = _parse_deadline(payload.end_date)
    if payload.status is not None:
        if payload.status not in VALID_CONTRACT_STATUSES:
            raise HTTPException(status_code=400, detail=f"Invalid status. Must be one of: {', '.join(sorted(VALID_CONTRACT_STATUSES))}")
        contract.status = payload.status
    if payload.performance_security is not None:
        contract.performance_security = payload.performance_security.strip()
    if payload.notes is not None:
        contract.notes = payload.notes

    contract.updated_at = datetime.utcnow()
    db.commit()
    return _contract_response(contract)


@app.delete("/contracts/{contract_id}")
def delete_contract(
    contract_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    contract = _get_org_contract(contract_id, membership, db)
    if not _can_modify_contract(contract, membership):
        raise HTTPException(status_code=403, detail="You can only delete contracts you created.")

    db.delete(contract)
    db.commit()
    return {"detail": "Contract deleted."}


# ---------------------------------------------------------------------------
# AI Proposal Wizard — generate personalized proposal with company KB
# ---------------------------------------------------------------------------

@app.post("/tenders/{tender_id}/generate-proposal")
@limiter.limit("5/minute")
async def generate_proposal(
    request: Request,
    tender_id: int,
    payload: ProposalWizardRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only generate proposals for tenders you created.")

    org = _get_org(membership.organization_id, db)

    # Build tender analysis dict from stored fields
    tender_analysis = {
        "summary": tender.summary,
        "eligibility": tender.eligibility,
        "financial_requirements": tender.financial_requirements,
        "required_documents": tender.required_documents,
        "compliance_matrix": tender.compliance_matrix,
        "risk_analysis": tender.risk_analysis,
    }

    # Merge knowledge base with basic profile
    kb = _get_knowledge_base(current_user)
    if not kb.get("company_name"):
        kb["company_name"] = org.name or ""
    if not kb.get("contact_name"):
        kb["contact_name"] = current_user.contact_name or ""
    if not kb.get("phone"):
        kb["phone"] = current_user.phone or ""
    if not kb.get("address"):
        kb["address"] = current_user.address or ""

    wizard_data = payload.model_dump()
    tender_id_copy = tender.id
    organization_id = membership.organization_id
    language = payload.language

    async def generate():
        yield _sse({"type": "progress", "stage": "generating", "message": "AI is writing your personalized proposal..."})

        loop = asyncio.get_running_loop()
        queue: asyncio.Queue = asyncio.Queue()

        def produce():
            try:
                for chunk in stream_personalized_proposal(tender_analysis, kb, wizard_data, language):
                    asyncio.run_coroutine_threadsafe(queue.put(("chunk", chunk)), loop).result()
            except Exception as exc:
                asyncio.run_coroutine_threadsafe(queue.put(("error", str(exc))), loop).result()
            finally:
                asyncio.run_coroutine_threadsafe(queue.put(("done", None)), loop).result()

        thread = threading.Thread(target=produce, daemon=True)
        thread.start()

        accumulated: list[str] = []
        while True:
            kind, data = await queue.get()
            if kind == "chunk":
                accumulated.append(data)
                yield _sse({"type": "chunk", "text": data})
            elif kind == "error":
                yield _sse({"type": "error", "detail": data})
                thread.join(timeout=5)
                return
            elif kind == "done":
                break

        thread.join(timeout=5)
        yield _sse({"type": "progress", "stage": "saving", "message": "Saving personalized proposal..."})

        full_text = "".join(accumulated)

        # Save to tender
        new_db = SessionLocal()
        try:
            t = new_db.query(Tender).filter(
                Tender.id == tender_id_copy, Tender.organization_id == organization_id
            ).first()
            if t:
                t.personalized_proposal = full_text
                new_db.commit()
        finally:
            new_db.close()

        yield _sse({"type": "done", "proposal": full_text})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ---------------------------------------------------------------------------
# AI Bid Strategy Advisor
# ---------------------------------------------------------------------------

@app.post("/tenders/{tender_id}/bid-strategy")
@limiter.limit("5/minute")
async def generate_bid_strategy(
    request: Request,
    tender_id: int,
    language: str = Form("english"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    tender = db.query(Tender).filter(
        Tender.id == tender_id, Tender.organization_id == membership.organization_id
    ).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found.")
    if not _can_modify_tender(tender, membership):
        raise HTTPException(status_code=403, detail="You can only generate bid strategy for tenders you created.")

    org = _get_org(membership.organization_id, db)

    tender_analysis = {
        "summary": tender.summary,
        "eligibility": tender.eligibility,
        "financial_requirements": tender.financial_requirements,
        "required_documents": tender.required_documents,
        "compliance_matrix": tender.compliance_matrix,
        "risk_analysis": tender.risk_analysis,
        "bid_recommendation": tender.bid_recommendation,
    }

    kb = _get_knowledge_base(current_user)
    kb.setdefault("company_name", org.name or "")
    kb.setdefault("contact_name", current_user.contact_name or "")
    kb.setdefault("phone", current_user.phone or "")
    kb.setdefault("address", current_user.address or "")

    tender_id_copy = tender.id
    organization_id = membership.organization_id

    async def generate():
        yield _sse({"type": "progress", "stage": "analyzing", "message": "AI generating bid strategy, compliance analysis, and price intelligence..."})

        loop = asyncio.get_running_loop()
        queue: asyncio.Queue = asyncio.Queue()

        def produce():
            try:
                for chunk in stream_bid_strategy(tender_analysis, kb, language):
                    asyncio.run_coroutine_threadsafe(queue.put(("chunk", chunk)), loop).result()
            except Exception as exc:
                asyncio.run_coroutine_threadsafe(queue.put(("error", str(exc))), loop).result()
            finally:
                asyncio.run_coroutine_threadsafe(queue.put(("done", None)), loop).result()

        thread = threading.Thread(target=produce, daemon=True)
        thread.start()

        accumulated: list[str] = []
        while True:
            kind, data = await queue.get()
            if kind == "chunk":
                accumulated.append(data)
                yield _sse({"type": "chunk", "text": data})
            elif kind == "error":
                yield _sse({"type": "error", "detail": data})
                thread.join(timeout=5)
                return
            elif kind == "done":
                break

        thread.join(timeout=5)
        full_text = "".join(accumulated)

        new_db = SessionLocal()
        try:
            t = new_db.query(Tender).filter(
                Tender.id == tender_id_copy, Tender.organization_id == organization_id
            ).first()
            if t:
                t.bid_strategy = full_text
                new_db.commit()
        finally:
            new_db.close()

        yield _sse({"type": "done"})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ---------------------------------------------------------------------------
# AI Document Validator
# ---------------------------------------------------------------------------

@app.post("/documents/validate")
@limiter.limit("20/minute")
async def validate_doc(
    request: Request,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    content = await file.read()
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large. Maximum is 15 MB.")

    ext = os.path.splitext(file.filename or "").lower()[1]
    mime = file.content_type or ""

    allowed_exts = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".pdf"}
    if mime not in {"image/jpeg", "image/jpg", "image/png", "image/webp", "image/gif", "application/pdf"} and ext not in allowed_exts:
        raise HTTPException(
            status_code=400,
            detail="Only image files (JPG/PNG/WEBP/GIF) and PDF documents are supported.",
        )

    if ext in {".jpg", ".jpeg"} and not mime.startswith("image/"):
        mime = "image/jpeg"
    elif ext == ".png" and not mime.startswith("image/"):
        mime = "image/png"
    elif ext == ".pdf":
        mime = "application/pdf"

    result = validate_document(content, mime, file.filename or "document")
    return result


# ---------------------------------------------------------------------------
# Calendar & Notifications
# ---------------------------------------------------------------------------

def _computed_reminders(db: Session, current_user: User, membership: OrgMembership) -> list[dict]:
    """Deadline/high-score/cert/contract-expiry reminders, recomputed on every
    call (not persisted — there's nothing meaningful to mark 'read' on a fact
    like 'this deadline is in 3 days', it's just always true until it isn't)."""
    now = datetime.utcnow()
    reminders: list[dict] = []

    upcoming = (
        db.query(Tender)
        .filter(
            Tender.organization_id == membership.organization_id,
            Tender.deadline.isnot(None),
            Tender.deadline >= now,
            Tender.deadline <= now + timedelta(days=14),
        )
        .order_by(Tender.deadline.asc())
        .all()
    )
    for t in upcoming:
        days_left = (t.deadline - now).days
        urgency = "critical" if days_left <= 2 else "warning" if days_left <= 5 else "info"
        reminders.append({
            "id": None, "persisted": False, "read_at": None,
            "type": "deadline", "entity_type": "tender", "entity_id": t.id,
            "title": t.title, "deadline": t.deadline, "days_left": days_left,
            "urgency": urgency,
            "message": (
                f"{'URGENT: ' if urgency == 'critical' else ''}"
                f"Tender \"{t.title}\" deadline in {days_left} day{'s' if days_left != 1 else ''}."
            ),
        })

    high_score_reviewing = (
        db.query(Tender)
        .filter(
            Tender.organization_id == membership.organization_id,
            Tender.bid_status == "reviewing",
            Tender.bid_score >= 70,
        )
        .order_by(Tender.bid_score.desc())
        .limit(5)
        .all()
    )
    for t in high_score_reviewing:
        if any(r["entity_id"] == t.id and r["type"] == "deadline" for r in reminders):
            continue
        reminders.append({
            "id": None, "persisted": False, "read_at": None,
            "type": "high_score", "entity_type": "tender", "entity_id": t.id,
            "title": t.title, "deadline": t.deadline,
            "days_left": (t.deadline - now).days if t.deadline else None,
            "urgency": "info",
            "message": (
                f"High-opportunity tender \"{t.title}\" (score {t.bid_score}/100) "
                "is still under review. Consider submitting."
            ),
        })

    contracts_expiring = (
        db.query(Contract)
        .filter(
            Contract.organization_id == membership.organization_id,
            Contract.status == "active",
            Contract.end_date.isnot(None),
            Contract.end_date >= now,
            Contract.end_date <= now + timedelta(days=30),
        )
        .order_by(Contract.end_date.asc())
        .all()
    )
    for c in contracts_expiring:
        days_left = (c.end_date - now).days
        urgency = "critical" if days_left <= 7 else "warning" if days_left <= 14 else "info"
        reminders.append({
            "id": None, "persisted": False, "read_at": None,
            "type": "contract_expiry", "entity_type": "contract", "entity_id": c.id,
            "title": c.title, "deadline": c.end_date, "days_left": days_left,
            "urgency": urgency,
            "message": f"Contract \"{c.title}\" ends in {days_left} day{'s' if days_left != 1 else ''}.",
        })

    try:
        kb = _get_knowledge_base(current_user)
        for cert in kb.get("certifications", []):
            exp_str = cert.get("expiry_date") or cert.get("expiry") or ""
            if not exp_str:
                continue
            try:
                exp_dt = datetime.fromisoformat(exp_str[:10])
            except ValueError:
                continue
            days_until_exp = (exp_dt - now).days
            if -30 <= days_until_exp <= 60:
                urgency = "critical" if days_until_exp < 0 else ("warning" if days_until_exp <= 14 else "info")
                name = cert.get("name") or cert.get("cert_name") or "Certificate"
                reminders.append({
                    "id": None, "persisted": False, "read_at": None,
                    "type": "cert_expiry", "entity_type": None, "entity_id": None,
                    "title": name, "deadline": exp_dt, "days_left": days_until_exp,
                    "urgency": urgency,
                    "message": (
                        f"{'EXPIRED: ' if days_until_exp < 0 else ''}"
                        f"Certificate \"{name}\" "
                        f"{'expired' if days_until_exp < 0 else 'expires'} "
                        f"{exp_dt.strftime('%d %b %Y')}."
                    ),
                })
    except Exception:
        pass

    return reminders


def _notification_response(n: Notification) -> dict:
    return {
        "id": n.id, "persisted": True,
        "type": n.type, "entity_type": n.entity_type, "entity_id": n.entity_id,
        "title": n.title, "message": n.message or "", "urgency": n.urgency or "info",
        "read_at": n.read_at, "created_at": n.created_at,
    }


@app.get("/notifications")
def list_notifications(
    unread_only: bool = False,
    limit: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    query = db.query(Notification).filter(
        Notification.organization_id == membership.organization_id,
        Notification.user_id == membership.user_id,
    )
    if unread_only:
        query = query.filter(Notification.read_at.is_(None))
    persisted = query.order_by(Notification.id.desc()).limit(max(1, min(limit, 200))).all()

    combined = [_notification_response(n) for n in persisted] + _computed_reminders(db, current_user, membership)
    urgency_order = {"critical": 0, "warning": 1, "info": 2}

    def sort_key(r: dict):
        unread = 0 if (not r["persisted"] or r["read_at"] is None) else 1
        urgency_rank = urgency_order.get(r["urgency"], 9)
        ts = r.get("created_at") or r.get("deadline") or datetime.utcnow()
        return (unread, urgency_rank, -ts.timestamp())

    combined.sort(key=sort_key)

    unread_count = sum(1 for r in combined if not r["persisted"] or r["read_at"] is None)
    return {"notifications": combined, "count": len(combined), "unread_count": unread_count}


@app.post("/notifications/{notification_id}/read")
def mark_notification_read(
    notification_id: int,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    notification = db.query(Notification).filter(
        Notification.id == notification_id,
        Notification.organization_id == membership.organization_id,
        Notification.user_id == membership.user_id,
    ).first()
    if not notification:
        raise HTTPException(status_code=404, detail="Notification not found.")
    notification.read_at = datetime.utcnow()
    db.commit()
    return _notification_response(notification)


@app.post("/notifications/read-all")
def mark_all_notifications_read(
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    now = datetime.utcnow()
    db.query(Notification).filter(
        Notification.organization_id == membership.organization_id,
        Notification.user_id == membership.user_id,
        Notification.read_at.is_(None),
    ).update({"read_at": now})
    db.commit()
    return {"detail": "All notifications marked read."}


@app.get("/calendar")
def get_calendar(
    from_: str | None = Query(default=None, alias="from"),
    to: str | None = None,
    db: Session = Depends(get_db),
    membership: OrgMembership = Depends(get_current_membership),
):
    now = datetime.utcnow()
    range_start = _parse_deadline(from_) or (now - timedelta(days=30))
    range_end = _parse_deadline(to) or (now + timedelta(days=180))

    events: list[dict] = []

    tenders = db.query(Tender).filter(
        Tender.organization_id == membership.organization_id,
        Tender.deadline.isnot(None),
        Tender.deadline >= range_start,
        Tender.deadline <= range_end,
    ).all()
    for t in tenders:
        events.append({
            "date": t.deadline, "type": "tender_deadline",
            "entity_type": "tender", "entity_id": t.id, "title": t.title,
            "urgency": "critical" if (t.deadline - now).days <= 2 else "warning" if (t.deadline - now).days <= 5 else "info",
        })

    contracts = db.query(Contract).filter(
        Contract.organization_id == membership.organization_id,
        Contract.end_date.isnot(None),
        Contract.end_date >= range_start,
        Contract.end_date <= range_end,
    ).all()
    for c in contracts:
        events.append({
            "date": c.end_date, "type": "contract_end",
            "entity_type": "contract", "entity_id": c.id, "title": c.title,
            "urgency": "warning" if (c.end_date - now).days <= 14 else "info",
        })

    tasks = db.query(TaskModel).filter(
        TaskModel.organization_id == membership.organization_id,
        TaskModel.due_date.isnot(None),
        TaskModel.due_date >= range_start,
        TaskModel.due_date <= range_end,
        TaskModel.status.notin_(("done", "cancelled")),
    ).all()
    for tk in tasks:
        events.append({
            "date": tk.due_date, "type": "task_due",
            "entity_type": tk.entity_type, "entity_id": tk.entity_id, "title": tk.title,
            "urgency": "warning" if (tk.due_date - now).days <= 2 else "info",
        })

    events.sort(key=lambda e: e["date"])
    return {"events": events}


# ---------------------------------------------------------------------------
# AI Tender Discovery
# ---------------------------------------------------------------------------

_discovery_state: dict = {"running": False, "last_run": None, "count": 0}


def _sync_discovered(db: Session, notices: list[dict]) -> int:
    new_count = 0
    for n in notices:
        existing = db.query(DiscoveredTender).filter(
            DiscoveredTender.external_id == n["external_id"]
        ).first()
        if existing:
            continue
        dt = DiscoveredTender(
            source=n.get("source", ""),
            external_id=n["external_id"],
            title=n.get("title", "")[:500],
            description=(n.get("description") or "")[:800],
            category=(n.get("category") or "")[:300],
            deadline=n.get("deadline"),
            estimated_value=n.get("estimated_value", "")[:200],
            url=(n.get("url") or "")[:500],
            country=n.get("country", "Bangladesh"),
        )
        db.add(dt)
        new_count += 1
    db.commit()
    return new_count


@app.get("/discover")
def list_discovered(
    source: str = "",
    search: str = "",
    limit: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    query = db.query(DiscoveredTender)
    if source.strip():
        query = query.filter(DiscoveredTender.source == source.strip())
    if search.strip():
        term = f"%{search.strip()}%"
        query = query.filter(
            or_(DiscoveredTender.title.ilike(term), DiscoveredTender.description.ilike(term))
        )
    items = query.order_by(DiscoveredTender.discovered_at.desc()).limit(max(1, min(limit, 200))).all()
    return {
        "tenders": [
            {
                "id": d.id,
                "source": d.source,
                "external_id": d.external_id,
                "title": d.title,
                "description": (d.description or "")[:300],
                "category": d.category,
                "deadline": d.deadline,
                "estimated_value": d.estimated_value,
                "url": d.url,
                "country": d.country,
                "discovered_at": d.discovered_at,
            }
            for d in items
        ],
        "total": len(items),
        "state": _discovery_state,
    }


@app.post("/discover/refresh")
@limiter.limit("3/minute")
async def refresh_discovery(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if _discovery_state["running"]:
        return {"detail": "Discovery already running.", "state": _discovery_state}

    _discovery_state["running"] = True

    def _run():
        try:
            notices = run_all_scrapers()
            fresh_db = SessionLocal()
            try:
                added = _sync_discovered(fresh_db, notices)
                _discovery_state["count"] = added
                _discovery_state["last_run"] = datetime.utcnow().isoformat()
            finally:
                fresh_db.close()
        except Exception as exc:
            import logging
            logging.getLogger(__name__).warning("Discovery refresh failed: %s", exc)
        finally:
            _discovery_state["running"] = False

    threading.Thread(target=_run, daemon=True).start()
    return {"detail": "Discovery refresh started.", "state": _discovery_state}


@app.post("/discover/{discovered_id}/import")
def import_discovered_tender(
    discovered_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    membership: OrgMembership = Depends(get_current_membership),
):
    dt = db.query(DiscoveredTender).filter(DiscoveredTender.id == discovered_id).first()
    if not dt:
        raise HTTPException(status_code=404, detail="Discovered tender not found.")

    summary = (
        f"Discovered from {dt.source}.\n\n"
        f"Category: {dt.category or 'N/A'}\n"
        f"Estimated Value: {dt.estimated_value or 'N/A'}\n"
        f"Country: {dt.country}\n"
        f"Source URL: {dt.url or 'N/A'}\n\n"
        f"{dt.description or ''}"
    ).strip()

    tender = Tender(
        user_id=current_user.id,
        organization_id=membership.organization_id,
        title=dt.title,
        language="english",
        status="completed",
        file_name="",
        file_type="",
        file_size=0,
        deadline=dt.deadline,
        bid_status="reviewing",
        original_text=summary,
        summary=summary,
    )
    db.add(tender)
    db.commit()
    db.refresh(tender)
    return get_tender_response(tender)
